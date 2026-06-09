# -*- coding: utf-8 -*-
"""
Tier 3：跨樣區 ANOVA / 混合模型 統計分析（接 cinnamon_analysis.py 之 _inc.pkl）
輸出：output/03_ANOVA統計分析.xlsx、output/03_ANOVA統計分析報告.docx、output/figs/*.png
模型（user 拍板）：
  反應變數＝年均增量 ΔD/yr、ΔH/yr 與 相對生長率 RGR_D；成熟/幼齡分開。
  幼齡：主分析＝混合模型（處理固定 修剪×施肥、樣區隨機）；交叉驗證＝樣區均值 2×2 ANOVA（樣區為複製，n=3）。
  成熟：僅 1 樣區/處理 → 樹級 2×2 ANOVA（偽複製，僅供描述/探索）。
"""
import os, math, warnings
import numpy as np
import pandas as pd
from scipy import stats
import statsmodels.api as sm
import statsmodels.formula.api as smf
from statsmodels.stats.anova import anova_lm
from statsmodels.stats.multicomp import pairwise_tukeyhsd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
warnings.filterwarnings("ignore")

import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(BASE, "output")
FIGS = os.path.join(OUT, "figs")
os.makedirs(FIGS, exist_ok=True)
inc = pd.read_pickle(os.path.join(OUT, "_inc.pkl"))

RESP = [("dD_annual", "年均直徑生長量 ΔD (cm/yr)"), ("dH_annual", "年均樹高生長量 ΔH (m/yr)"), ("RGR_D", "直徑相對生長率 RGR_D (/yr)")]
TRT_ORDER = ["C0", "P1", "F150", "P1F150"]

# matplotlib 中文（Windows 微軟正黑體；失敗則退英文標籤）
try:
    plt.rcParams["font.sans-serif"] = ["Microsoft JhengHei", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False
except Exception:
    pass

xl_tables = {}   # sheetname -> df


def two_way(data, resp):
    m = smf.ols(f"{resp} ~ C(pruning)*C(fertilizer)", data=data).fit()
    aov = anova_lm(m, typ=2).reset_index().rename(columns={"index": "來源", "sum_sq": "SS", "df": "df", "F": "F", "PR(>F)": "p"})
    # 假設檢定
    sh = stats.shapiro(m.resid)
    grp = [g[resp].values for _, g in data.groupby("treatment")]
    lev = stats.levene(*grp)
    return aov, m, dict(shapiro_W=round(sh.statistic, 3), shapiro_p=round(sh.pvalue, 4),
                        levene_F=round(lev.statistic, 3), levene_p=round(lev.pvalue, 4))


def tukey_tbl(data, resp):
    r = pairwise_tukeyhsd(data[resp].values, data["treatment"].values)
    d = pd.DataFrame(r.summary().data[1:], columns=r.summary().data[0])
    return d


def mixed(data, resp):
    md = smf.mixedlm(f"{resp} ~ C(pruning)*C(fertilizer)", data, groups=data["site"])
    mf = md.fit(reml=True, method="lbfgs")
    t = pd.DataFrame({"項": mf.params.index, "係數": mf.params.values,
                      "標準誤": mf.bse.reindex(mf.params.index).values, "z": mf.tvalues.values, "p": mf.pvalues.values})
    t = t[~t["項"].str.contains("Group Var")]
    t[["係數", "標準誤", "z", "p"]] = t[["係數", "標準誤", "z", "p"]].astype(float).round(4)
    return t, mf


def boxfig(data, resp, title, fname):
    fig, ax = plt.subplots(figsize=(5.2, 3.4))
    order = [t for t in TRT_ORDER if t in data["treatment"].unique()]
    vals = [data[data["treatment"] == t][resp].values for t in order]
    ax.boxplot(vals, labels=order, showmeans=True)
    ax.axhline(0, color="#999", lw=0.8, ls="--")
    ax.set_title(title, fontsize=10)
    ax.set_ylabel(resp)
    ax.set_xlabel("處理")
    fig.tight_layout()
    p = os.path.join(FIGS, fname)
    fig.savefig(p, dpi=140)
    plt.close(fig)
    return p


# ===================== 跑分析 =====================
results = {"mature": {}, "juvenile": {}}
figpaths = []

# 成熟林：樹級 2×2（偽複製）
mat = inc[inc.group == "mature"].copy()
for resp, lab in RESP:
    aov, model, assump = two_way(mat, resp)
    tuk = tukey_tbl(mat, resp)
    results["mature"][resp] = dict(aov=aov, assump=assump, tukey=tuk, label=lab)
    xl_tables[f"成熟_{resp}_ANOVA"] = aov
    xl_tables[f"成熟_{resp}_Tukey"] = tuk
    figpaths.append((("mature", resp), boxfig(mat, resp, f"成熟林 115121：{lab}", f"mature_{resp}.png")))

# 幼齡林：混合模型 + 樣區均值 2×2
juv = inc[inc.group == "juvenile"].copy()
juv_plotmean = (juv.groupby(["plot", "site", "treatment", "pruning", "fertilizer"])
                .agg(dD_annual=("dD_annual", "mean"), dH_annual=("dH_annual", "mean"), RGR_D=("RGR_D", "mean"))
                .reset_index())
for resp, lab in RESP:
    mt, mf = mixed(juv, resp)
    aov_pm, model_pm, assump_pm = two_way(juv_plotmean, resp)
    tuk_pm = tukey_tbl(juv_plotmean, resp)
    results["juvenile"][resp] = dict(mixed=mt, aov_pm=aov_pm, assump_pm=assump_pm, tukey_pm=tuk_pm, label=lab)
    xl_tables[f"幼齡_{resp}_混合模型"] = mt
    xl_tables[f"幼齡_{resp}_樣區均值ANOVA"] = aov_pm
    xl_tables[f"幼齡_{resp}_Tukey樣區均值"] = tuk_pm
    figpaths.append((("juvenile", resp), boxfig(juv, resp, f"幼齡林（3區）：{lab}", f"juv_{resp}.png")))

# ===================== Excel =====================
xlpath = os.path.join(OUT, "03_ANOVA統計分析.xlsx")
with pd.ExcelWriter(xlpath, engine="openpyxl") as xw:
    for name, d in xl_tables.items():
        d.to_excel(xw, sheet_name=name[:31], index=False)
print("  ✓ Excel:", xlpath)

# ===================== Word 報告 =====================
CJK = "DFKai-SB"
doc = docx.Document()
# 預設字型
st = doc.styles["Normal"]
st.font.name = CJK
st.font.size = Pt(12)
st.element.rPr.rFonts.set(qn("w:eastAsia"), CJK)


def setfont(run, sz=12, bold=False, italic=False):
    run.font.name = CJK
    run.font.size = Pt(sz)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CJK)


def H(txt, sz=15):
    p = doc.add_paragraph()
    r = p.add_run(txt); setfont(r, sz, True)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    return p


def body(txt, indent=True):
    p = doc.add_paragraph()
    r = p.add_run(txt); setfont(r, 12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    return p


def li(txt):
    p = doc.add_paragraph()
    r = p.add_run(txt); setfont(r, 12)
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.space_after = Pt(3)
    return p


def sci_body(pre, sci, post):
    p = doc.add_paragraph()
    r1 = p.add_run(pre); setfont(r1, 12)
    r2 = p.add_run(sci); setfont(r2, 12, italic=True)
    r3 = p.add_run(post); setfont(r3, 12)
    p.paragraph_format.first_line_indent = Pt(24)
    return p


def add_table(df, ndp=4):
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    for j, c in enumerate(df.columns):
        cell = t.rows[0].cells[j]; cell.text = ""
        setfont(cell.paragraphs[0].add_run(str(c)), 10, True)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(df.columns):
            v = row[c]
            if isinstance(v, float):
                v = "<0.0001" if (c in ("p", "p-adj") and v < 0.0001) else round(v, ndp)
            cells[j].text = ""
            setfont(cells[j].paragraphs[0].add_run(str(v)), 9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_fig(path, width_cm=11):
    doc.add_picture(path, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


# 封面/標題
ttl = doc.add_paragraph(); ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
setfont(ttl.add_run("土肉桂修枝矮化及施肥對生長影響之研究"), 18, True)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
setfont(sub.add_run("處理間生長反應變異數分析報告"), 15, True)
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
setfont(meta.add_run("資料期別：第一期(114Q3)～第三期(115Q1)　分析日期：中華民國 115 年 6 月".replace("　", "  ")), 11)
doc.add_paragraph()

H("壹、分析方法")
sci_body("本研究樹種為土肉桂（", "Cinnamomum osmophloeum", "）。試驗採二因子析因配置加對照，二因子為修剪（0/1）與施肥（0/150 公克），共 4 種處理（C0 對照、P1 修剪、F150 施肥、P1F150 修剪加施肥），配置於 4 處樣區。")
body("反應變數採每株自第一期至第三期之年均生長量計算：年均直徑生長量 ΔD（公分／年）、年均樹高生長量 ΔH（公尺／年），及直徑相對生長率 RGR_D（以自然對數差除以期距年數）。僅納入兩端期皆存活且具有效徑高之個體。")
li("一、成熟林（樣區 115121，量胸徑）與幼齡林（樣區 117218、811、8110，量地徑）因量測部位與尺度不同，分開分析。")
li("二、幼齡林：主分析採線性混合模型（修剪與施肥為固定效應、樣區為隨機效應），以處理樣區之重複（每處理 3 樣區）為推論基礎；另以樣區均值進行 2×2 析因變異數分析交叉驗證，並以 Tukey HSD 進行事後多重比較。")
li("三、成熟林：每處理僅 1 樣區，無樣區層重複，以樹為單位之 2×2 析因變異數分析僅供描述與探索（存在偽複製，結論不宜外推）。")
li("四、各模型均檢定殘差常態性（Shapiro–Wilk）與組間變異齊一性（Levene）。分析以 Python（statsmodels、scipy）執行，顯著水準 α＝0.05。")

# 樣本概況
H("貳、樣本概況")
nrow = (inc.groupby(["group"]).agg(樣本樹數=("treeNum", "count"), 平均期距年=("years", "mean")).round(3).reset_index())
nrow["group"] = nrow["group"].map({"mature": "成熟林(115121)", "juvenile": "幼齡林(3區)"})
nrow = nrow.rename(columns={"group": "立地"})
add_table(nrow, 3)

# 成熟林
H("參、成熟林分析（樣區 115121，胸徑）")
body("註：本節每處理僅 1 樣區、以樹為單位分析，存在偽複製，且現有期距僅約半年、量測噪音與多莖二次平均徑之莖組變動使年均增量波動偏大，下列結果僅供探索參考。", indent=True)
for resp, lab in RESP:
    r = results["mature"][resp]
    H(f"一、{lab}", 13)
    body("2×2 析因變異數分析（型 II）：", indent=False)
    add_table(r["aov"])
    a = r["assump"]
    body(f"假設檢定：Shapiro–Wilk W＝{a['shapiro_W']}（p＝{a['shapiro_p']}）；Levene F＝{a['levene_F']}（p＝{a['levene_p']}）。", indent=False)
    add_fig(dict(figpaths)[("mature", resp)])

# 幼齡林
H("肆、幼齡林分析（樣區 117218、811、8110，地徑）")
for resp, lab in RESP:
    r = results["juvenile"][resp]
    H(f"一、{lab}", 13)
    body("（一）線性混合模型（固定：修剪×施肥；隨機：樣區）固定效應：", indent=False)
    add_table(r["mixed"])
    body("（二）樣區均值 2×2 析因變異數分析（交叉驗證，每處理 n＝3 樣區）：", indent=False)
    add_table(r["aov_pm"])
    a = r["assump_pm"]
    body(f"假設檢定（樣區均值）：Shapiro–Wilk W＝{a['shapiro_W']}（p＝{a['shapiro_p']}）；Levene F＝{a['levene_F']}（p＝{a['levene_p']}）。", indent=False)
    add_fig(dict(figpaths)[("juvenile", resp)])

# 結論與限制
H("伍、結論與限制")
li("一、解讀提醒：現有資料僅 3 期、期距約半年，年均增量受季節性量測誤差放大；土肉桂多莖以二次平均徑彙整，各期所測莖數／身分變動會使單株徑值跳動，造成部分樣區出現負成長之假象。故本報告之處理效果檢定屬初步，待第四期起時間序列延長後再行確認。")
li("二、樹高量測於半年期距內出現大幅正負變動，研判為測高之量測噪音；建議第四期加強樹高量測一致性（固定測點、同一儀器與人員），並考慮以基面積增量作為多莖整株生長之輔助指標。")
li("三、統計推論層級：幼齡林以樣區為處理重複（n＝3）方屬正確推論單位，混合模型同此精神；成熟林僅 1 樣區／處理，無法就處理效果作正式推論，需增設重複或改以時間序列趨勢評估。")
li("四、後續建議：第四期完成後，將期距延長之增量重跑本管線；必要時加入「期別」為重複量測之時間因子（重複量測混合模型），以提升統計檢定力。")

docpath = os.path.join(OUT, "03_ANOVA統計分析報告.docx")
doc.save(docpath)
print("  ✓ Word:", docpath)
print("完成。")
