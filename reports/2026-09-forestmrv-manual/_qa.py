# -*- coding: utf-8 -*-
import sys, io
try:
    sys.stdout.reconfigure(encoding='utf-8')
except Exception:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import zipfile, re
from pathlib import Path
from docx import Document

HERE = Path(__file__).parent
DOCX = HERE / "操作手冊-ForestMRV智慧森林監測平臺.docx"

# 1) 原始 XML 層級檢查（U+3000、項目符號、斜體、圖）
z = zipfile.ZipFile(str(DOCX))
xml = z.read('word/document.xml').decode('utf-8')

u3000 = xml.count('　')
# 項目符號：numbering 內 bullet 或 document 內 numPr（清單）— python-docx 我們未用內建 list，故應為 0
numpr = xml.count('<w:numPr')
italic = xml.count('<w:i/>') + xml.count('<w:i ')
# 圖片
media = [n for n in z.namelist() if n.startswith('word/media/')]

print("docx 存在：", DOCX.exists(), " 大小：%.2f MB" % (DOCX.stat().st_size/1e6))
print("U+3000 全形空格：", u3000, "（須為 0）")
print("numPr 內建項目符號/編號清單：", numpr, "（我們用手打編號，應為 0）")
print("斜體 run（學名等）：", italic)
print("嵌入圖片數：", len(media))

# 2) python-docx 層級：段落數、標題、待補佔位
doc = Document(str(DOCX))
paras = doc.paragraphs
n_para = len(paras)
headings = [p.text for p in paras if p.style.name.startswith('Heading')]
placeholders = sum(1 for t in doc.tables for row in t.rows for c in row.cells
                   for p in c.paragraphs if '待補截圖' in p.text)
print("段落總數：", n_para, " 表格數：", len(doc.tables))
print("Heading 段落數：", len(headings))
print("待補佔位框：", placeholders)

# 3) 抓任何殘留全形空格的上下文
if u3000:
    for m in re.finditer('.{0,20}　.{0,20}', xml):
        print("  U+3000 脈絡：", m.group(0))

# 4) 章標題列出
print("\n--- 章節標題（Heading 1/2 前 40）---")
for p in paras:
    if p.style.name in ('Heading 1', 'Heading 2'):
        print(" ", p.style.name, "|", p.text)

print("\n--- 待拍清單 ---")
cl = HERE / "_需拍攝清單.md"
if cl.exists():
    print(cl.read_text(encoding='utf-8'))
