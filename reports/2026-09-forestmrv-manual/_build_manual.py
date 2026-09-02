# -*- coding: utf-8 -*-
"""
產生：ForestMRV 智慧森林監測平臺 — 完整操作手冊（對應 prod v2.11.95）

依 CLAUDE.md 全域排版規則：
  - 5 級標號：壹、 / 一、 / (一) / 1. / (1)
  - 段落首行縮排 2 字
  - 無項目符號（一律編號清單）
  - 無全形空格（U+3000）
  - 學名斜體
  - CJK 字型：PMingLiU（中文） + Times New Roman（西文）

圖片策略：
  - 沿用兩套既有截圖（不複製，直接引用原始夾）：
      IMG_GEN = 2026-05-mid-demo/_build/images   （5/20 通用系統圖，M/P/H 系列）
      IMG_CIN = 2026-06-cinnamon-demo/images      （6/6 土肉桂修枝圖，S 系列 + 封面底圖）
  - 新功能（v2.11.69→95）無截圖者 → add_placeholder 留佔位框，並自動彙整成
    「需拍攝清單」印到 console + 寫入 _需拍攝清單.md（交 user 補拍）。

跑法：python _build_manual.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
OUT = HERE / "操作手冊-ForestMRV智慧森林監測平臺.docx"
IMG_GEN = HERE.parent / "2026-05-mid-demo" / "_build" / "images"
IMG_CIN = HERE.parent / "2026-06-cinnamon-demo" / "images"
IMG_NEW = HERE / "images"          # user 補拍的新圖放這

CJK_FONT = "PMingLiU"
ASCII_FONT = "Times New Roman"

CHECKLIST = []   # 收集待補圖，build 後輸出清單


def _imgdir(src):
    return {'gen': IMG_GEN, 'cin': IMG_CIN, 'new': IMG_NEW}[src]


# ───────────────────────── 樣式 helper ─────────────────────────
def set_cjk_font(run, name=CJK_FONT, ascii_name=ASCII_FONT, size=12):
    run.font.name = ascii_name
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), ascii_name)
    rFonts.set(qn('w:hAnsi'), ascii_name)


def add_para(doc, text, size=12, bold=False, indent=True, color=None, align=None, after_pt=4, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after_pt)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.85)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    set_cjk_font(run, size=size)
    return p


def add_para_runs(doc, runs, indent=True, after_pt=4):
    """混排：runs = [(text, {'italic':True,'bold':True}), ...]，供學名斜體用。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after_pt)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.85)
    for text, opt in runs:
        r = p.add_run(text)
        r.italic = bool(opt.get('italic'))
        r.bold = bool(opt.get('bold'))
        set_cjk_font(r, size=opt.get('size', 12))
    return p


def add_heading(doc, text, level, after_pt=6):
    sizes = {1: 16, 2: 14, 3: 13, 4: 12, 5: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(after_pt)
    if level >= 3:
        p.paragraph_format.first_line_indent = Cm(0.85)
    run = p.add_run(text)
    run.bold = True
    set_cjk_font(run, size=sizes.get(level, 12))
    # 標記為 Heading 樣式以利目次（保留自訂字型/字級）
    if level <= 3:
        try:
            p.style = doc.styles['Heading %d' % level]
            run.bold = True
            set_cjk_font(run, size=sizes.get(level, 12))
        except Exception:
            pass
    return p


def add_title(doc, text, size=20):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    set_cjk_font(run, size=size)
    return p


def add_subtitle(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_cjk_font(run, size=size)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def add_image(doc, filename, src='gen', caption=None, width_inches=5.5):
    img_path = _imgdir(src) / filename
    if not img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("［圖檔遺失：%s／%s］" % (src, filename))
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.italic = True
        set_cjk_font(run, size=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        crun = cp.add_run(caption)
        crun.italic = True
        crun.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        set_cjk_font(crun, size=10)


PH_SEEN = set()   # 已使用的圖號，防重複


def add_placeholder(doc, code, caption, shot_desc, width_inches=5.5):
    """新功能圖：圖號由呼叫端明確指定（穩定識別），不隨文件位置浮動。
       若 images/Nxx.png 已存在 → 自動嵌入該截圖（不再列入待拍清單）；
       否則 → 畫灰底佔位框並登錄到待拍清單。

       ⚠ 舊版此處是「依文件出現順序自動編號」，忽略傳入的 code。只要在文件中間
       插入一個新小節，後面每一張既有截圖都會被綁到別人的圖名上（v2.11.95 改版
       時實際踩到，11 張全錯位）。圖號即截圖檔名，必須是穩定識別、不可位置相關。"""
    if code in PH_SEEN:
        raise ValueError("佔位圖號重複：%s（圖號即檔名，必須唯一）" % code)
    PH_SEEN.add(code)
    img_path = IMG_NEW / (code + ".png")
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        crun = cp.add_run(caption)
        crun.italic = True
        crun.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        set_cjk_font(crun, size=10)
        return
    CHECKLIST.append((code, caption, shot_desc))
    # 佔位框（單格表格，灰底虛線感）
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.width = Cm(13)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'F3F4F6'); tcPr.append(shd)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run("［待補截圖 %s］" % code)
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)
    set_cjk_font(r1, size=11)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(shot_desc)
    r2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    set_cjk_font(r2, size=9)
    # 圖名
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(10)
    crun = cp.add_run(caption)
    crun.italic = True
    crun.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    set_cjk_font(crun, size=10)


def add_table_with_header(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        set_cjk_font(run, size=11)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'DCE6F1'); tcPr.append(shd)
    for r_i, row in enumerate(rows, 1):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i].cells[c_i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_cjk_font(run, size=10)
    if col_widths:
        for col, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[col].width = Cm(w)
    return t


def add_callout(doc, text, kind='tip'):
    icons = {'tip': '💡 小提示', 'warn': '⚠️ 注意', 'rule': '📜 法規說明', 'tech': '🔧 系統設計說明'}
    colors = {'tip': RGBColor(0x15, 0x80, 0x3D), 'warn': RGBColor(0xC2, 0x41, 0x0C),
              'rule': RGBColor(0x1D, 0x4E, 0xD8), 'tech': RGBColor(0x6B, 0x21, 0xA8)}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(0.5)
    icon_run = p.add_run(icons.get(kind, '💡') + ' ')
    icon_run.bold = True
    icon_run.font.color.rgb = colors.get(kind, RGBColor(0x15, 0x80, 0x3D))
    set_cjk_font(icon_run, size=11)
    body_run = p.add_run(text)
    set_cjk_font(body_run, size=11)


def add_toc(doc):
    """插入 Word 自動目次欄位（開檔時按 F9 或允許更新即生成頁碼）。"""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "（開啟本檔後，於目次上按右鍵→更新功能變數，即可產生頁碼）"
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instr); run._r.append(fldChar2)
    run._r.append(t); run._r.append(fldChar3)
    set_cjk_font(run, size=11)


def page_break(doc):
    doc.add_page_break()


# ═══════════════════════ 開始建文件 ═══════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ── 封面 ──
for _ in range(3):
    doc.add_paragraph()
add_title(doc, "ForestMRV 智慧森林監測平臺", size=24)
add_title(doc, "操作手冊", size=22)
doc.add_paragraph()
add_subtitle(doc, "測量・報告・查證（MRV）森林資源監測與碳匯會計系統", size=13)
add_subtitle(doc, "對應正式版本 v2.11.95", size=12)
add_subtitle(doc, "中華民國 115 年 9 月（第二版）", size=12)
doc.add_paragraph()
add_subtitle(doc, "https://forestry-rs-monitor.web.app", size=11)
page_break(doc)

# ── 目次 ──
add_heading(doc, "目次", 1)
add_toc(doc)
page_break(doc)


# ═══════════════════════ 章節內容 ═══════════════════════
import _content as _c     # 章節內容拆到 _content.py，保持本檔可讀
_c.build(globals())

# ── 存檔 ──
doc.save(str(OUT))
print("OK ->", OUT.name)

# ── 輸出待拍清單 ──
if CHECKLIST:
    lines = ["# ForestMRV 操作手冊 — 需補拍截圖清單", "",
             "> 自動產生自 _build_manual.py。手冊內對應位置已留佔位框。",
             "> 拍好後請依「建議檔名」存到 images/ 夾，再重跑 python _build_manual.py 即自動嵌入。",
             "", "| 圖號 | 建議檔名 | 圖名 / 內容 | 拍攝指引 |", "|---|---|---|---|"]
    for code, caption, desc in CHECKLIST:
        fn = code + ".png"
        lines.append("| %s | %s | %s | %s |" % (code, fn, caption, desc))
    (HERE / "_需拍攝清單.md").write_text("\n".join(lines), encoding="utf-8")
    print("待拍清單 ->", "_需拍攝清單.md（共 %d 張）" % len(CHECKLIST))
