"""
產生：02-操作手冊-土肉桂專區葉片收穫監測會計系統.docx
為 6/6 土肉桂專區系統說明會準備（對應 prod v2.11.69）。

依 CLAUDE.md 全域排版規則：
  - 5 級標號：壹、 / 一、 / (一) / 1. / (1)
  - 段落首行縮排 2 字
  - 無項目符號（編號清單）
  - 無全形空格（U+3000）
  - CJK 字型：PMingLiU（中文） + Times New Roman（西文）

對象：林業保育署臺中分署承辦、土肉桂專區合作社代表、契約林農、計畫團隊內部。
範圍：純土肉桂專區葉片收穫監測會計系統功能（不含樣區調查、立木普查等
其他模組，那些功能屬於 ForestMRV 主系統其他章節）。

跑法：python _build_manual.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
OUT = HERE / "02-操作手冊-土肉桂專區葉片收穫監測會計系統.docx"
IMG = HERE / "images"

# ===== 樣式設定 =====
CJK_FONT = "PMingLiU"
ASCII_FONT = "Times New Roman"


def set_cjk_font(run, name=CJK_FONT, ascii_name=ASCII_FONT, size=12):
    run.font.name = ascii_name
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), ascii_name)
    rFonts.set(qn('w:hAnsi'), ascii_name)


def add_para(doc, text, size=12, bold=False, indent=True, color=None, align=None, after_pt=4, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after_pt)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.85)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    set_cjk_font(run, size=size)
    return p


def add_heading(doc, text, level, after_pt=6):
    sizes = {1: 16, 2: 14, 3: 13, 4: 12, 5: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(after_pt)
    if level >= 3:
        p.paragraph_format.first_line_indent = Cm(0.85)
    run = p.add_run(text)
    run.bold = True
    set_cjk_font(run, size=sizes.get(level, 12))
    return p


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    set_cjk_font(run, size=20)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(text)
    set_cjk_font(run, size=12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_image(doc, filename, caption=None, width_inches=5.5):
    """嵌入截圖 + 圖名（置中、無首行縮排）"""
    img_path = IMG / filename
    if not img_path.exists():
        # 占位符（圖檔遺失時提示）
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"［圖檔遺失：{filename}］")
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run.italic = True
        set_cjk_font(run, size=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        crun = cp.add_run(caption)
        crun.italic = True
        crun.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        set_cjk_font(crun, size=10)


def add_table_with_header(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        set_cjk_font(run, size=11)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'DCE6F1')
        tcPr.append(shd)
    for r_i, row in enumerate(rows, 1):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i].cells[c_i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_cjk_font(run, size=10)
    if col_widths:
        for col, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[col].width = Cm(w)


def add_callout(doc, text, kind='tip'):
    """加突顯框（小提示／注意／警告）"""
    icons = {'tip': '💡 小提示', 'warn': '⚠️ 注意', 'rule': '📜 法規說明', 'tech': '🔧 系統設計說明'}
    colors = {'tip': RGBColor(0x15, 0x80, 0x3D), 'warn': RGBColor(0xC2, 0x41, 0x0C),
              'rule': RGBColor(0x1D, 0x4E, 0xD8), 'tech': RGBColor(0x6B, 0x21, 0xA8)}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(0.5)
    icon_run = p.add_run(icons.get(kind, '💡') + ' ')
    icon_run.bold = True
    icon_run.font.color.rgb = colors.get(kind, RGBColor(0x15, 0x80, 0x3D))
    set_cjk_font(icon_run, size=11)
    body_run = p.add_run(text)
    set_cjk_font(body_run, size=11)


# ===== 開始建文件 =====
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

add_title(doc, "土肉桂專區葉片收穫監測會計系統")
add_title(doc, "操作手冊")
add_subtitle(doc, "ForestMRV 智慧森林監測平臺 │ v2.11.69 │ 2026 年 5 月 30 日")

# ─────────────────────────────────────────────────────
# 壹、系統概論
# ─────────────────────────────────────────────────────
add_heading(doc, "壹、系統概論", 1)

add_heading(doc, "一、系統定位", 2)
add_para(doc,
    "本系統名為「土肉桂專區葉片收穫監測會計系統」，內建於 ForestMRV 智慧森林監測"
    "平臺，依循 IPCC LULUCF 與 MRV（測量、報告、查證）原則設計。針對土肉桂專區，"
    "本系統承接以下行政流程：")
add_para(doc, "1. 林農申請修枝採取副產物（葉片）；")
add_para(doc, "2. 林業保育署臺中分署核准並核發法定許可文號；")
add_para(doc, "3. 林農回報實際葉片採收量；")
add_para(doc, "4. 結案後資料移交合作社進行共同銷售彙整。")

add_para(doc,
    "本系統採行政管理導向設計，與林分樣區調查、立木普查、QAQC 等專案類型分屬不同"
    "模組，操作介面已簡化為純行政流程所需之七個分頁（地圖／修枝申請／修枝審核／"
    "葉片採收回報及結案／葉片採收彙整／匯出／設定）。")

add_heading(doc, "二、系統使用對象與角色", 2)
add_para(doc,
    "本系統設計五個操作角色，各角色職責分明、權限以 Firestore Security Rules "
    "強制執行：")
role_def = [
    ["角色（系統名）", "對應人員", "核心職責", "本手冊章節"],
    ["系統管理員（admin）", "計畫團隊", "開案、設邊界、人員指派、QAQC 監督", "伍"],
    ["計畫主持人（pi）", "計畫團隊指派／合作社主席（移交後）", "方法學設定、樣區規格、邀請成員", "伍（部分共用）"],
    ["林農（surveyor）", "契約承租林農", "修枝申請、實際葉片採收量回報、結案", "貳"],
    ["審核人員（harvest_authority）", "林業保育署臺中分署承辦", "核准／駁回／要求補件、核發法定許可文號", "參"],
    ["合作社（coop）", "合作社主席", "唯讀彙整檢視、共同銷售匯出 Excel", "肆"],
]
add_table_with_header(doc, role_def[0], role_def[1:], col_widths=[3.0, 4.5, 5.5, 1.8])

add_callout(doc,
    "現階段（計畫執行期）系統管理員與計畫主持人皆由計畫團隊負責人擔任。計畫結束移交後，"
    "系統管理員仍由計畫團隊保留以維護資料完整性；計畫主持人將移交予合作社主席。"
    "林農、審核人員、合作社角色於兩階段皆保持一致。",
    kind='rule')

add_heading(doc, "三、系統入口與登入", 2)
add_para(doc,
    "本系統為 Progressive Web App（PWA），無需安裝任何軟體，於任何 Chrome 或 "
    "Edge 瀏覽器開啟以下網址即可使用：")
add_para(doc, "https://forestry-rs-monitor.web.app", indent=False, bold=True)
add_para(doc,
    "首次連線後即可於登入畫面選擇「使用 Google 登入」或「Email/密碼登入」其中"
    "一種方式。建議使用 Google 帳號登入，避免另記密碼。")

add_image(doc, "S01-登入頁.png",
          caption="圖 1-1 系統登入頁面（兩種登入方式）", width_inches=4.0)

add_callout(doc,
    "若需新建系統帳號，請聯絡計畫團隊（系統管理員）開設並指派角色。直接於登入頁"
    "「Email/密碼登入」展開「註冊」功能僅供測試用途，新註冊帳號預設無任何專案"
    "權限，須由系統管理員另外加入土肉桂專案。",
    kind='warn')

add_heading(doc, "四、登入後的主畫面與專案首頁", 2)
add_para(doc,
    "登入成功後，系統會顯示「我的專案」清單。視您的角色不同，可看到的專案數量會"
    "有所差異：admin 看得到全部專案；pi/surveyor/harvest_authority/coop 僅看得到"
    "被指派的專案。對於本說明會的所有與會者，皆應能看到「土肉桂專區葉片收穫監測、"
    "會計系統」這張卡片。")

add_image(doc, "S02-我的專案.png",
          caption="圖 1-2 我的專案清單（admin 視角，可看到全部專案）", width_inches=5.5)

add_para(doc,
    "點選「土肉桂專區葉片收穫監測、會計系統」卡片進入專案首頁。專案首頁上方為"
    "頁首區（含返回、編輯專案、專案名稱、編號），中段為七個分頁列，下方為當前"
    "分頁內容區。預設停留於「地圖」分頁，立即顯示 115 個作業單元的邊界疊圖。")

add_image(doc, "S03-土肉桂專案首頁.png",
          caption="圖 1-3 土肉桂專案首頁（地圖分頁、含 115 個作業單元 label）", width_inches=6.0)

add_callout(doc,
    "土肉桂專區為純行政專案（不含樣區調查），系統已自動隱藏「樣區／設計／儀表板」"
    "三個與本業務無關的分頁，介面更乾淨。其他類型的 ForestMRV 專案（如森林經營"
    "計畫、碳匯專案）會顯示完整分頁，不適用本手冊範圍。",
    kind='tech')

# ─────────────────────────────────────────────────────
# 貳、林農端操作（修枝申請 → 採收回報 → 結案）
# ─────────────────────────────────────────────────────
add_heading(doc, "貳、林農端操作（修枝申請 → 採收回報 → 結案）", 1)

add_para(doc,
    "本章節說明契約林農於本系統內之完整操作流程，包含修枝申請、查看核准結果、"
    "列印申請函/許可單、回報實際葉片採收量、以及結案。系統內林農之系統角色名稱"
    "為 surveyor。")

add_heading(doc, "一、修枝申請業務流程概覽", 2)
add_para(doc,
    "依林業及自然保育署相關規定，林農於土肉桂專區內進行修枝作業並採取副產物"
    "（葉片）前，須先向臺中分署提出申請、待核准後方可進行。本系統將此行政流程"
    "電子化，並於各環節保留紙本公文格式以利存檔。完整業務流程如下：")
flow_steps = [
    ["步驟", "動作", "操作者", "本系統內對應功能"],
    ["1", "提出修枝申請（含林地位置、面積、株數、期間）", "林農", "🌿 修枝申請 → ＋ 修枝申請"],
    ["2", "列印申請函（公文格式）寄送或交付分署", "林農", "📄 申請公文稿"],
    ["3", "分署審查並核准／要求補件／駁回", "分署承辦", "📋 修枝審核（見參章）"],
    ["4", "核准後系統自動配發法定許可文號", "系統", "（自動）"],
    ["5", "列印採取許可單（公文格式）作為作業期間隨身證明", "林農", "📋 採取許可單"],
    ["6", "於核准效期內進行修枝作業並採收葉片", "林農", "（線下作業）"],
    ["7", "回報實際葉片採收量（可分批多次）", "林農", "🌾 葉片採收回報及結案"],
    ["8", "全數回報完畢後結案", "林農", "✅ 回報完畢並結案"],
    ["9", "結案後資料供合作社彙整、共同銷售", "合作社", "📊 葉片採收彙整（見肆章）"],
]
add_table_with_header(doc, flow_steps[0], flow_steps[1:], col_widths=[1.5, 7.5, 2.0, 4.5])

add_heading(doc, "二、進入修枝申請分頁", 2)
add_para(doc,
    "於專案首頁上方分頁列點選「🌿 修枝申請」分頁，即可看到本人歷來所有修枝申請"
    "案件清單。各案件以卡片形式呈現，依狀態（草稿／已送出．待審／已核准／修枝"
    "作業中／已結案／已駁回）分區顯示。")

add_image(doc, "S04-修枝申請清單.png",
          caption="圖 2-1 修枝申請分頁（含三筆已結案案件範例）", width_inches=6.0)

add_para(doc,
    "頁面上方「土肉桂修枝申請」標題下方有業務說明文字，提醒林農：修枝受林業保育"
    "署臺中分署管控，須先提出申請、核准後始得修枝及採取枝葉。送出後待分署審核，"
    "核准前可存草稿並編輯；核准後請至「🌾 葉片採收回報及結案」分頁填報葉片採收量。")

add_para(doc,
    "右上角綠色「＋ 修枝申請」按鈕為新增申請的入口。點選後系統會開啟新增申請"
    "表單（modal 形式）。")

add_heading(doc, "三、新增修枝申請（六步驟）", 2)

add_heading(doc, "(一) 開啟新增申請表單", 3)
add_para(doc,
    "點選「＋ 修枝申請」按鈕後，系統顯示「修枝申請（土肉桂）」表單視窗。表單"
    "由上至下依序為：作業位置 picker、迷你地圖預覽、申請人姓名、聯絡方式、林班/"
    "地號、申請修枝面積、修枝株數、修枝方式、修枝起迄日、備註、儲存按鈕。")

add_image(doc, "S05-新增申請-空白表單.png",
          caption="圖 2-2 新增修枝申請表單（空白狀態、預設先看到 picker 與地圖）", width_inches=4.0)

add_callout(doc,
    "本表單依「先選作業位置 → 各欄位自動帶入」的順序設計。一旦完成上方三層 picker "
    "選擇，下方多個欄位（申請人姓名、林班/地號、面積、株數）會自動填入，林農只需"
    "補上「聯絡方式」與「修枝起迄日」兩個必填欄位即可送出。", kind='tip')

add_heading(doc, "(二) 第一層 picker：選專區", 3)
add_para(doc,
    "於「📍 作業位置（從專案邊界帶入）」區塊內，第一個下拉「1. 專區」列出本"
    "專案邊界涵蓋之專區。土肉桂專區共有兩個地理區域：")
add_para(doc, "1. 橫流溪（45 個作業單元）")
add_para(doc, "2. 烏石坑（70 個作業單元）")
add_para(doc, "請依您的契約林地所在區域選擇對應專區。")

add_heading(doc, "(三) 第二層 picker：選承租人", 3)
add_para(doc,
    "選定專區後，第二個下拉「2. 承租人」會自動列出該專區所有契約承租人姓名"
    "（字典序排列）。請從清單中選取本人姓名。")
add_callout(doc,
    "選定承租人的同時，上方「申請人姓名」欄位會自動填入該承租人姓名（系統假設"
    "「承租人＝申請人」為原則）。若實際申請人非承租人本人（如代理人），可於表單"
    "送出前手動修改此欄位。", kind='tip')

add_heading(doc, "(四) 第三層 picker：選作業單元", 3)
add_para(doc,
    "選定承租人後，第三個下拉「3. 作業單元」會列出該承租人名下所有作業單元，"
    "每筆顯示作業單元編號（如「14_1_1」或「117-107-01」）、契約面積（ha）、"
    "契約樹種摘要。請選取本次申請欲進行修枝之作業單元。")

add_image(doc, "S06-Picker三層連動完成.png",
          caption="圖 2-3 picker 三層連動完成（mini-map 紅色高亮目標作業單元、申請人姓名自動帶入）",
          width_inches=4.0)

add_para(doc,
    "三層選定完成後，picker 下方會顯示「✓ 已帶入下方林班/地號欄位」摘要框，"
    "列出七欄結構化資料：專區、承租人、作業單元、契約樹種、契約面積、工作站/"
    "事業區、契約書號。同時下方迷你地圖（Leaflet 嵌入）會自動紅色高亮所選"
    "作業單元並 zoom 至該位置（maxZoom 18）。")

add_callout(doc,
    "若您 picker 選錯（例如選錯承租人或作業單元），可直接於 picker 下拉重新選擇，"
    "系統會即時更新下方欄位與地圖高亮。所有變動僅在表單內，未送出前不會寫入系統。",
    kind='tip')

add_heading(doc, "(五) 補填必填欄位（聯絡方式、修枝起迄日）", 3)
add_para(doc,
    "完成上方 picker 後，請補填以下三個必填欄位（皆標示紅色星號 *）：")
add_para(doc,
    "1. 聯絡方式（電話／email）：用於分署承辦人於審查中需聯繫補件時使用，"
    "亦會列印於申請公文稿與許可單上。")
add_para(doc,
    "2. 修枝起日：本次申請欲開始進行修枝作業之日期。")
add_para(doc,
    "3. 修枝迄日：本次申請預計完成修枝作業之日期（須晚於起日）。系統會檢查"
    "起迄日順序，若起日晚於迄日將擋下送出。")

add_image(doc, "S07-申請表單下半部.png",
          caption="圖 2-4 申請表單下半部（林班/地號、面積、株數已自動帶入；修枝起迄日待填）",
          width_inches=4.5)

add_para(doc,
    "其餘欄位多由 picker 自動帶入或系統預設，林農可視實際狀況微調：")
add_para(doc,
    "1. 林班/地號：picker 帶入；如有差異可手動編輯覆蓋。")
add_para(doc,
    "2. 申請修枝面積（ha）：picker 帶入契約面積；如本次申請僅部分區塊"
    "可手動修改。")
add_para(doc,
    "3. 土肉桂修枝株數：系統依面積 × 1800 株/ha（合作社契約假設密度）"
    "自動計算填入，修改面積時即時連動更新；可手動覆蓋。")
add_para(doc,
    "4. 修枝方式：預設「修枝」（亦可選「其他」）。")
add_para(doc,
    "5. 備註：選填，可記載特殊作業考量或補充說明。")

add_heading(doc, "(六) 儲存草稿或送出申請", 3)
add_para(doc,
    "表單最下方有兩個按鈕：")
add_para(doc,
    "1. 💾 儲存草稿（淺灰按鈕）：將申請存為草稿狀態，可日後繼續編輯或補件。"
    "草稿不會配發發文字號，亦不會進入分署待審清單。")
add_para(doc,
    "2. 📤 送出申請（深綠按鈕）：將申請正式送出，系統會即時配發發文字號"
    "「大雪山林業合作社（修）字第 NNN 號」（三位流水號自動編號），並進入分署"
    "「📋 修枝審核」分頁待審清單。")

add_callout(doc,
    "發文字號採 Firestore runTransaction 原子遞增機制，確保每筆送出都拿到"
    "唯一連號。已配發過的案件（如改卷、再送出）不會重新編號，避免破壞紙本"
    "公文連續性。",
    kind='tech')

add_heading(doc, "四、列印申請公文稿", 2)
add_para(doc,
    "送出申請後，於申請清單該案件卡片上點選「📄 申請公文稿」按鈕。系統會於"
    "新分頁開啟正式公文「函」格式之申請函，可直接以瀏覽器列印功能輸出紙本或"
    "另存 PDF 備存。")

add_image(doc, "S08-申請公文稿-第1頁.png",
          caption="圖 2-5 申請公文稿第 1 頁（公文體 4 點說明 + 表列 8 欄申請明細）",
          width_inches=5.0)

add_para(doc,
    "申請公文稿第 1 頁採正式公文「函」格式，包含：")
add_para(doc, "1. 受文者：林業及自然保育署臺中分署")
add_para(doc, "2. 發文日期：以送出當日為準")
add_para(doc, "3. 發文字號：「大雪山林業合作社（修）字第 NNN 號」")
add_para(doc, "4. 速別、密等、附件三欄")
add_para(doc, "5. 主旨：申請於下列林地進行土肉桂修枝及枝葉採取乙案，請核准")
add_para(doc, "6. 說明（四點正式公文體）")
add_para(doc, "7. 表列資料：申請人/聯絡方式/林班地號/修枝面積/株數/修枝方式/期間/備註")

add_image(doc, "S09-申請公文稿-第2頁SVG.png",
          caption="圖 2-6 申請公文稿第 2 頁（附圖：作業位置示意圖 + 申請人簽章 + 分署收件欄）",
          width_inches=5.0)

add_para(doc,
    "申請公文稿第 2 頁附件包含：")
add_para(doc, "1. 作業位置示意圖：自動產生之 SVG 圖形，紅色實心區域為本次申請作業單元，"
              "灰色為鄰近作業單元（context），右上角為地理北方箭頭。")
add_para(doc, "2. 圖下方 meta 資訊：作業單元代碼、專區、承租人。")
add_para(doc, "3. 申請人簽章欄、身分證/統一編號欄、申請日期欄（請林農手寫簽章補齊）。")
add_para(doc, "4. 分署收件用方框（收文日期、收文文號、承辦人）—留白由分署填寫。")

add_callout(doc,
    "申請公文稿之發文字號、表列資料、附圖等皆即時由線上記錄產生，紙本內容與"
    "線上紀錄完全一致、不會漂移。若送出後線上資料有變更（例如要求補件後修改），"
    "重新列印一份新的公文稿即可。", kind='tech')

add_callout(doc,
    "印出申請公文稿並親簽蓋章後，請依分署規定方式（紙本郵寄、現場遞交、或掃描"
    "回傳 email）送達分署承辦人。線上送出狀態同時通知分署，但**仍以紙本公文"
    "為法定文書**。",
    kind='rule')

add_heading(doc, "五、查看核准結果與列印採取許可單", 2)
add_para(doc,
    "分署承辦人於系統內審核後（核准／要求補件／駁回），林農端會即時看到狀態"
    "變更：")
add_para(doc,
    "1. 已核准：卡片狀態徽章變綠色「已核准」，新出現「📜 採取許可單」按鈕，"
    "並顯示效期。系統同時配發法定許可文號「林保中-土肉桂修枝-{民國年}-{NNN}」。")
add_para(doc,
    "2. 要求補件：狀態徽章變橙色「要求補件」，卡片下方顯示審核意見。林農可點"
    "「✏️ 編輯」修改後重新送出。")
add_para(doc,
    "3. 駁回：狀態徽章變紅色「已駁回」，卡片下方顯示駁回理由。本案件流程結束。")

add_para(doc,
    "於已核准案件卡片上點選「📜 採取許可單」按鈕，系統會於新分頁開啟採取許可單"
    "公文格式，可列印作為作業期間隨身證明。")

add_image(doc, "S12-採取許可單-列印版.png",
          caption="圖 2-7 採取許可單列印版（公文格式、含 v2.11.69 新增宣告）",
          width_inches=5.0)

add_para(doc,
    "採取許可單之公文 box 內容包含：")
add_para(doc, "1. 許可文號：「林保中-土肉桂修枝-115-NNN」")
add_para(doc, "2. 申請人姓名與聯絡方式")
add_para(doc, "3. 林地（林班/地號）與面積")
add_para(doc, "4. 修枝方式與土肉桂修枝株數")
add_para(doc, "5. 效期（起日 ～ 迄日）")
add_para(doc, "6. 審核附註（若分署有填）")
add_para(doc, "7. **底部一句宣告**（dashed border-top 區隔）：本許可只許可於上述林地進行"
              "土肉桂修枝作業；所採取葉片重量以申請人於 ForestMRV 線上系統實際回報為準、"
              "不設預估上限。")

add_callout(doc,
    "依 v2.11.69 業務拍板：本系統許可不預估／不限制葉片採收量上限。許可只許可"
    "您進行修枝作業，所採取之葉片重量請於修枝完成後於系統內據實回報，系統不會"
    "比對核准量與實際量、亦無「達成率」「超量警告」等概念。",
    kind='rule')

add_heading(doc, "六、回報實際葉片採收量", 2)
add_para(doc,
    "於核准效期內完成修枝作業並採收葉片後，請於系統內回報實際採收重量。回報"
    "為法定義務，未回報不得結案。")

add_para(doc,
    "切換至「🌾 葉片採收回報及結案」分頁，可看到待回報之核准案件。每筆案件"
    "卡片顯示許可文號、林地、效期、已回報累計量。")

add_image(doc, "S14-葉片採收回報及結案.png",
          caption="圖 2-8 葉片採收回報及結案分頁（含三筆已結案案件範例）",
          width_inches=6.0)

add_para(doc,
    "點選案件卡片上「＋ 填報葉片採收量」按鈕，開啟登錄表單。表單需填寫：")
add_para(doc,
    "1. 採收日期（必填）：本批次採收的實際日期。")
add_para(doc,
    "2. 實際鮮葉重 (kg)（必填）：本次採收的鮮葉重量。")
add_para(doc,
    "3. 乾燥後重 (kg)：選填，若已乾燥處理可一併記錄。")
add_para(doc,
    "4. 含水率 (%)：選填，乾燥處理後之含水率（0～100）。")
add_para(doc,
    "5. 批次標示：選填，自定義批次代號便於分批管理。")
add_para(doc,
    "6. 備註：選填。")

add_para(doc,
    "送出後系統會即時：")
add_para(doc,
    "1. 將本筆登錄寫入該許可單之 logs 子集合。")
add_para(doc,
    "2. 重新加總所有 logs 寫入許可單 totalLogged_kg 欄位。")
add_para(doc,
    "3. 若為首次回報，自動推進案件狀態 approved → harvesting（修枝作業中）。")
add_para(doc,
    "4. 顯示 toast「✅ 已回報，累計 X kg」。")

add_callout(doc,
    "葉片採收回報可分批多次。例如同一許可文號下，您可於 5/30 採收 10 kg 回報"
    "一次、6/15 採收 8 kg 回報另一次，系統累計顯示 18 kg。每筆回報都會出現"
    "在卡片「已回報明細」表內，依採收日期排序。",
    kind='tip')

add_heading(doc, "七、回報完畢並結案", 2)
add_para(doc,
    "當您判斷本許可案件所有葉片採收已回報完畢、不再新增時，請按案件卡片上"
    "「✅ 回報完畢並結案」按鈕，系統會跳出確認對話框列出：")
add_para(doc, "1. 許可文號")
add_para(doc, "2. 累計回報葉片採收量 (kg)")
add_para(doc, "3. 結案聲明文字")

add_para(doc,
    "確認後案件狀態將推進至「已結案（completed）」，往後不能再填報葉片採收量，"
    "亦不能再編輯案件內容。結案後資料固定，供分署查核與合作社共同銷售彙整使用。")

add_callout(doc,
    "結案閘門：若您「累計回報葉片採收量」為 0（從未回報）即點結案，系統會擋下"
    "並提示「請先＋填報葉片採收量，回報完畢後才能結案」。此為法規要求之回報"
    "義務閘門，由 Firestore Security Rules 與客戶端雙重把關。",
    kind='warn')

# ─────────────────────────────────────────────────────
# 參、分署承辦端操作（審核 → 核准 → 法定許可文號）
# ─────────────────────────────────────────────────────
add_heading(doc, "參、分署承辦端操作（審核 → 核准 → 法定許可文號）", 1)

add_para(doc,
    "本章節說明林業及自然保育署臺中分署承辦人員於本系統內之操作流程。系統內"
    "分署承辦之系統角色名稱為 harvest_authority。承辦人之權限受 Firestore "
    "Security Rules 嚴格限制：只能對待審／審核中之案件進行核准／要求補件／"
    "駁回，不能變更林農申請內容、不能跨專案讀寫、不能讀其他角色的個人草稿。")

add_heading(doc, "一、進入修枝審核分頁", 2)
add_para(doc,
    "於專案首頁上方分頁列點選「📋 修枝審核」分頁。若有待審案件，分頁名稱右側"
    "會出現紅色數字徽章顯示待審筆數。")

add_image(doc, "S10-修枝審核清單.png",
          caption="圖 3-1 修枝審核分頁（含三筆已處理案件範例）",
          width_inches=6.0)

add_para(doc,
    "頁面標題下方有副標清楚說明承辦權限：「審核者僅能核准／要求補件／駁回，"
    "不可變更申請內容（權限由 Firestore 規則鎖定）」。下方頁面分兩區塊：")
add_para(doc, "1. 待審核（NN）：尚未審核之案件，依送出時間排序。")
add_para(doc, "2. 已處理（NN）：已核准／要求補件／駁回之歷史案件。")

add_para(doc,
    "已處理區之每張卡片顯示：申請人姓名 + 已核准/已駁回徽章 + 法定許可文號（已"
    "核准案才有）+ 林地 + 面積 + 株數 + 修枝期間 + 系統識別（picker meta）+ "
    "效期 + 「📋 採取許可單」按鈕（可預覽該案件已配發的許可單）。")

add_heading(doc, "二、開啟待審案件並選擇審核動作", 2)
add_para(doc,
    "於「待審核」區塊找到欲審核之案件，案件卡片下方有三個按鈕：")
add_para(doc, "1. ✅ 核准（綠色）：通過審核並配發法定許可文號。")
add_para(doc, "2. 📋 要求補件（橙色）：請林農補件後再送。")
add_para(doc, "3. ✕ 駁回（紅色）：不予核准，本案件流程結束。")

add_para(doc,
    "點選對應按鈕後，系統會開啟對應審核 modal。三種動作的 modal 設計如下節說明。")

add_heading(doc, "三、核准修枝申請（核心流程）", 2)
add_para(doc,
    "點選「✅ 核准」後，系統開啟「核准修枝申請」modal。此 modal 結構與內容如下：")

add_image(doc, "S11-核准Modal-v2.11.69.png",
          caption="圖 3-2 核准修枝申請 modal（v2.11.69 新版，無「核准採收量」欄位）",
          width_inches=4.5)

add_para(doc,
    "Modal 內容由上至下依序為：")
add_para(doc, "1. 申請人資訊摘要（read-only）：申請人姓名／林地／修枝面積／株數一行帶過。")
add_para(doc, "2. 灰底提示框（v2.11.69 新增）：")
add_para(doc,
    "「本案核准後即許可進行修枝作業；所採取葉片重量由林農於 ForestMRV 線上系統"
    "實際回報，不設預估上限。」", indent=True)
add_para(doc, "3. 效期起日（預設為申請人填寫的修枝起日）。")
add_para(doc, "4. 效期迄日（預設為申請人填寫的修枝迄日）。")
add_para(doc, "5. 審核附註（選填）：可記載特殊核准條件、提醒事項等。")
add_para(doc, "6. ✅ 確認核准按鈕。")

add_callout(doc,
    "依 v2.11.69 業務拍板：本系統許可不預估／不限制葉片採收量上限。承辦人於"
    "核准時無需填寫「核准採收量」欄位，僅需確認效期即可。理由：土肉桂每塊地"
    "年齡不同、每株產出葉量差異大，事先預估上限反而誤導承辦與林農。系統採"
    "「事後實際回報」設計，林農回報量直接彙整供合作社共同銷售，無預估上限"
    "比對機制。",
    kind='rule')

add_para(doc,
    "確認核准後，系統會：")
add_para(doc, "1. 以 Firestore runTransaction 原子遞增 counters/harvestPermit 計數器。")
add_para(doc, "2. 配發法定許可文號「林保中-土肉桂修枝-{民國年}-{NNN}」三位流水號。")
add_para(doc, "3. 將申請狀態 submitted → approved，寫入效期與審核附註。")
add_para(doc, "4. 顯示 toast「✅ 已核准，許可文號 林保中-土肉桂修枝-115-NNN」。")
add_para(doc, "5. 該案件移出待審核區、進入已處理區，待審徽章 -1。")

add_callout(doc,
    "法定許可文號採 Firestore runTransaction 確保連號、不可重號。同一專案內所有"
    "承辦人核准的案件共用同一計數器，按核准順序配發。請勿在線上 console 直接"
    "修改 counters/harvestPermit 文件，避免破壞流水序。",
    kind='tech')

add_heading(doc, "四、要求補件", 2)
add_para(doc,
    "若申請內容有不足之處需林農補件後再送，請點「📋 要求補件」按鈕。Modal 內容："
    )
add_para(doc, "1. 申請人資訊摘要（read-only）。")
add_para(doc, "2. 理由（必填）：請具體說明需補哪些資料或修正哪些內容。")
add_para(doc, "3. ↩️ 退回補件按鈕。")

add_para(doc,
    "確認後系統將案件狀態 submitted → revision（補件中），並寫入 reviewComment。"
    "林農端會看到案件狀態徽章變橙色，可點「✏️ 編輯」依承辦指示修改後重新送出。")

add_heading(doc, "五、駁回", 2)
add_para(doc,
    "若申請內容不符規定、不予核准，請點「✕ 駁回」按鈕。Modal 內容同要求補件，"
    "理由必填。確認後系統將案件狀態 submitted → rejected，本案件流程結束、"
    "林農不能再編輯重送。")

add_callout(doc,
    "駁回為最終決定。若僅是申請內容需修正，請優先使用「要求補件」讓林農有機會"
    "改正；駁回應僅用於「申請事項本身不符規定」之情況（如非契約承租人申請、"
    "申請面積超出契約範圍等）。",
    kind='warn')

add_heading(doc, "六、查看已核准案件之採取許可單", 2)
add_para(doc,
    "於「已處理」區點選任一已核准案件卡片之「📋 採取許可單」按鈕，系統會顯示"
    "該案件之採取許可單預覽 modal，包含許可文號、申請人、林地、修枝方式、效期、"
    "宣告文字、葉片採收登錄表（依林農實際回報資料更新）、累計鮮葉採收量。"
    "點 modal 內「🖨️ 列印林產物採取許可單」按鈕可開啟列印頁面。")

# ─────────────────────────────────────────────────────
# 肆、合作社端操作（彙整 → 匯出 Excel）
# ─────────────────────────────────────────────────────
add_heading(doc, "肆、合作社端操作（彙整 → 匯出 Excel）", 1)

add_para(doc,
    "本章節說明土肉桂專區合作社於本系統內之操作流程。系統內合作社之系統角色"
    "名稱為 coop。合作社現階段為「唯讀觀察者」角色，可掌握各林農申請與葉片"
    "採收資訊、供共同銷售規劃，但無任何核准或編輯權限（權限由 Firestore "
    "規則天然鎖定）。未來移交後合作社主席將升級為計畫主持人（pi）角色，"
    "接手日常營運。")

add_heading(doc, "一、進入葉片採收彙整分頁", 2)
add_para(doc,
    "於專案首頁上方分頁列點選「📊 葉片採收彙整」分頁。本分頁僅顯示「已送出（含）"
    "起以後」的案件，亦即排除各林農的個人草稿（合作社不應看到他人未送出之草稿）。")

add_image(doc, "S13-葉片採收彙整.png",
          caption="圖 4-1 葉片採收彙整分頁（含三區塊：總覽 + 依林農彙整 + 各申請案）",
          width_inches=6.0)

add_para(doc,
    "本分頁分為三個區塊：")

add_heading(doc, "二、區塊一：專區總覽", 2)
add_para(doc,
    "顯示本專案各申請案件之總覽統計，包含：")
add_para(doc, "1. 申請案總數（已送出起）：例如「3 件」。")
add_para(doc, "2. 已核准/修枝作業中/結案案件數：例如「3 件」。")
add_para(doc,
    "3. 狀態分布：列出各狀態對應筆數，例如「已結案 3」。")
add_para(doc,
    "4. 已回報葉片採收量（事後實際累計）：所有已送出案件之 totalLogged_kg 加總，"
    "例如「45.0 kg」。此數字為合作社共同銷售規劃之主要依據。")

add_heading(doc, "三、區塊二：依林農葉片採收彙整（共同銷售）", 2)
add_para(doc,
    "本區塊以表格呈現各林農之葉片採收累計量，三欄結構：")
add_para(doc, "1. 林農：林農姓名（依採收量降序排列）。")
add_para(doc, "2. 案件數：該林農名下已送出申請案件總數。")
add_para(doc, "3. 已回報葉片採收量(kg)：該林農所有案件 totalLogged_kg 加總。")

add_para(doc,
    "本區塊主要用於共同銷售前彙整：合作社可依本表了解各林農預計交付之鮮葉重量、"
    "規劃通路與銷售策略。")

add_callout(doc,
    "「已回報葉片採收量」＝核准後實際葉片採收回報累計（修枝作業中／已結案），"
    "即可投入共同銷售之數量（鮮葉重）。系統不會比對「核准量 vs 實際量」、亦無"
    "達成率指標，所有彙整以實際回報為準。",
    kind='rule')

add_heading(doc, "四、區塊三：各申請案（唯讀）", 2)
add_para(doc,
    "本區塊列出各申請案之單筆資訊（合作社唯讀檢視），每張卡片顯示：")
add_para(doc, "1. 申請人姓名與狀態徽章。")
add_para(doc, "2. 法定許可文號（已核准案才有）。")
add_para(doc, "3. 林地（林班/地號）。")
add_para(doc, "4. 已回報葉片採收量 (kg)。")
add_para(doc, "5. 「📋 採取許可單」按鈕（可預覽該案件採取許可單）。")

add_heading(doc, "五、匯出 Excel 供共同銷售規劃", 2)
add_para(doc,
    "於專案首頁上方分頁列點選「匯出」分頁，可下載本專案完整資料為 Excel 檔案"
    "或多種 CSV 格式。合作社共同銷售規劃主要使用：")
add_para(doc, "1. 匯出 Excel（全部）：產出單一 .xlsx 檔，含本專案所有資料分頁。")
add_para(doc, "2. CSV: 收穫：單獨匯出葉片採收回報明細表。")

add_para(doc,
    "匯出之 Excel 可於 LibreOffice、Microsoft Excel 開啟編輯，用於後續銷售報表"
    "製作、與通路商議價、結算分潤等用途。")

# ─────────────────────────────────────────────────────
# 伍、系統管理員端操作（開案 → 邊界 → 成員 → 管理）
# ─────────────────────────────────────────────────────
add_heading(doc, "伍、系統管理員端操作", 1)

add_para(doc,
    "本章節說明系統管理員於本系統內之操作流程。系統內管理員之系統角色名稱為 "
    "admin（isSystemAdmin=true）。admin 為頂層角色，跨專案具最高權限，可開新"
    "專案、修改任何專案邊界、指派人員、跨狀態刪除案件等。")

add_heading(doc, "一、新建專案（開案）", 2)
add_para(doc,
    "於「我的專案」清單頁右上角點「＋ 新專案」按鈕，開啟新建專案表單。"
    "表單包含專案名稱、編號、類型、方法學等基本欄位。")

add_image(doc, "S15-開案表單.png",
          caption="圖 5-1 新建專案表單（含預設邊界一鍵載入功能）",
          width_inches=5.0)

add_para(doc,
    "對於土肉桂專區類型專案，開案時建議流程：")
add_para(doc, "1. 填寫專案名稱（如「土肉桂專區葉片收穫監測、會計系統」）與編號。")
add_para(doc, "2. 計畫類型選擇「行政管理」或「林產物採取監測」。")
add_para(doc,
    "3. 方法學與調查模組：勾選「admin_harvest」（林產物採取監測），其他樣區/"
    "立木/QAQC 等模組可全部不勾。")
add_para(doc,
    "4. 專案邊界（GeoJSON）：點「📥 預設邊界」下拉，選「土肉桂專區（橫流溪+"
    "烏石坑，合作社契約）」並點「載入」按鈕。系統會自動帶入合併邊界檔，含 115 個"
    "作業單元（45 + 70）與完整契約 metadata。")
add_para(doc, "5. 送出後新專案即建立完成。")

add_callout(doc,
    "土肉桂專區為「純行政專案」，不需要樣區調查、立木普查等模組。系統會依方法學"
    "勾選之模組自動隱藏不相關的分頁（樣區／設計／儀表板）與地圖頁的立木相關控制"
    "元件（著色依據／樣區邊界 checkbox），介面更乾淨。",
    kind='tech')

add_heading(doc, "二、邀請成員並指派角色", 2)
add_para(doc,
    "進入專案後，於上方分頁列點「設定」分頁，滾動至「成員管理」區塊。")

add_image(doc, "S16-設定分頁-成員管理.png",
          caption="圖 5-2 設定分頁成員管理區塊（含角色下拉與移除按鈕）",
          width_inches=5.5)

add_para(doc,
    "本區塊顯示專案目前所有成員，每列含成員姓名、email、角色下拉、移除按鈕（✕）。"
    "操作方式：")
add_para(doc,
    "1. 邀請成員：於下方「＋ 邀請成員」輸入框輸入欲邀請成員之 email，選擇對應"
    "角色（pi/surveyor/harvest_authority/coop/reviewer），按「邀請」按鈕。系統"
    "會自動把該成員加入專案，該成員下次登入即可看到此專案。")
add_para(doc,
    "2. 變更角色：點該成員列之角色下拉選單，改選新角色後系統即時生效。")
add_para(doc,
    "3. 移除成員：點該成員列右側「✕」按鈕，確認後該成員自此專案移除（其歷史"
    "資料保留，但無法再讀寫本專案）。")

add_callout(doc,
    "成員管理權限分級：admin 可加/移除任何成員（含其他 pi）；pi 可加/移除 "
    "surveyor、reviewer 與本身以外的 coop，但不能移除其他 pi（避免互踢無主）。"
    "所有角色都不能移除自己。",
    kind='warn')

add_heading(doc, "三、瀏覽地圖與作業單元 metadata", 2)
add_para(doc,
    "於專案首頁上方點「地圖」分頁，預設停留於此分頁。地圖會立即顯示本專案邊界"
    "與所有作業單元，每個作業單元以藍色邊框疊圖、上方有作業單元代碼之 tooltip "
    "label。")

add_image(doc, "S17-地圖頁115作業單元.png",
          caption="圖 5-3 地圖頁顯示作業單元代碼 label，點面跳出 7 欄結構化 metadata popup",
          width_inches=6.0)

add_para(doc,
    "操作互動：")
add_para(doc, "1. 點任一作業單元面（polygon）：跳出 popup 顯示 7 欄結構化 metadata（作業單元/"
              "專區/承租人/契約樹種/契約面積/工作站/契約書）。")
add_para(doc, "2. 滾輪縮放、拖曳平移：可探索整個專區範圍。")
add_para(doc, "3. 著色依據 / 疊加圖層控制：在純行政專案下已隱藏（因不適用），於含樣區之專案才會顯示。")

add_heading(doc, "四、刪除測試或誤送案件", 2)
add_para(doc,
    "本系統提供 admin/pi 跨狀態刪除案件功能，主要用途：")
add_para(doc, "1. demo 或教育訓練後一鍵清除測試資料。")
add_para(doc, "2. 處理林農誤送、重複申請、需作廢之案件。")
add_para(doc, "3. 過期未進入後續流程之案件清理。")

add_para(doc,
    "操作方式：於「修枝申請」或「修枝審核」分頁找到欲刪除之案件，點卡片上之"
    "「🗑️ 刪除申請案」紅框按鈕。系統會跳出確認對話框，列出將被刪除之資料：")

add_image(doc, "S18-PI刪除案件-確認對話框.png",
          caption="圖 5-4 刪除申請案確認對話框（admin 跨狀態可見）",
          width_inches=5.0)

add_para(doc,
    "確認對話框顯示內容：")
add_para(doc, "1. 申請人姓名")
add_para(doc, "2. 林地")
add_para(doc, "3. 案件狀態")
add_para(doc, "4. 發文字號（合作社編）")
add_para(doc, "5. 法定許可文號（分署編，已核准案才有，警示將同時失效）")
add_para(doc, "6. 累計回報量（警示子集合 logs 將一併刪除）")

add_callout(doc,
    "刪除為不可逆動作。對已核准案件，刪除會連同法定許可文號一併失效（counters 不"
    "回復）；對已結案案件，刪除會連同所有採收量回報明細一併移除。建議在刪除前"
    "對重要資料先匯出 Excel 備份。",
    kind='warn')

# ─────────────────────────────────────────────────────
# 陸、常見問題 FAQ
# ─────────────────────────────────────────────────────
add_heading(doc, "陸、常見問題（FAQ）", 1)

add_heading(doc, "一、登入與帳號相關", 2)

add_heading(doc, "(一) Q：忘記密碼怎麼辦？", 3)
add_para(doc,
    "A：建議優先使用「Google 登入」避免另記密碼。若使用 Email/密碼登入忘記密碼，"
    "請聯絡系統管理員（計畫團隊）重設。系統不開放自助密碼重設功能（避免社工攻擊）。")

add_heading(doc, "(二) Q：登入後看不到專案怎麼辦？", 3)
add_para(doc,
    "A：請確認您的角色已被系統管理員指派加入土肉桂專案。如尚未加入，請聯絡系統"
    "管理員 email 您要加入的專案名稱與您的登入 email。")

add_heading(doc, "(三) Q：可以在手機上使用嗎？", 3)
add_para(doc,
    "A：可以。本系統為 PWA（Progressive Web App），於 Chrome、Safari 等手機"
    "瀏覽器皆可使用。手機上可將網址加到主畫面，使用體驗近似 native app。")

add_heading(doc, "二、修枝申請相關（林農）", 2)

add_heading(doc, "(一) Q：picker 找不到我的承租人姓名怎麼辦？", 3)
add_para(doc,
    "A：請確認您選對「專區」（橫流溪 / 烏石坑）。若仍找不到，可能是合作社契約"
    "資料尚未更新至最新版本，請聯絡合作社或系統管理員確認契約名冊。")

add_heading(doc, "(二) Q：作業單元清單只有舊地號、找不到對應的新地號？", 3)
add_para(doc,
    "A：系統內作業單元代碼以合作社契約附圖為準（橫流溪用「14_1_1」格式、烏石坑"
    "用「117-107-01」格式）。若您手上的紙本資料用其他編號系統，請對照地理位置"
    "選擇對應之作業單元；如無法對應，請聯絡合作社協助。")

add_heading(doc, "(三) Q：申請已送出但發現填錯了，可以改嗎？", 3)
add_para(doc,
    "A：送出後狀態為「已送出．待審」時無法直接編輯。請聯絡分署承辦人「要求"
    "補件」退回，您即可編輯後重送。若分署已核准，編輯需與分署協調作廢重送。")

add_heading(doc, "(四) Q：株數一定要 1800/ha 嗎？", 3)
add_para(doc,
    "A：1800 株/ha 為合作社契約假設密度，系統依此預設值自動計算。實際密度可能"
    "因苗木損失、地形限制而有差異，可手動覆蓋此預設值改為實際估計株數。")

add_heading(doc, "三、葉片採收回報相關", 2)

add_heading(doc, "(一) Q：有沒有上限？採太多會被罰嗎？", 3)
add_para(doc,
    "A：依本系統 v2.11.69 業務拍板：許可只許可修枝作業、不預估／不限制葉片"
    "採收量上限。所採取葉片重量以您實際回報為準，系統不會比對核准量與實際量、"
    "亦無達成率指標。請依實際採收結果據實回報。若實際採收量遠高於預期，建議"
    "於備註欄說明原因供分署參考。")

add_heading(doc, "(二) Q：可以分批回報嗎？", 3)
add_para(doc,
    "A：可以。同一許可文號下，您可分批多次回報（如 5/30 採收 10 kg 回報一次、"
    "6/15 採收 8 kg 回報另一次），系統累計顯示。每筆回報都會出現在卡片「已回報"
    "明細」表內，依採收日期排序。")

add_heading(doc, "(三) Q：回報錯了可以改嗎？", 3)
add_para(doc,
    "A：目前 v2.11.69 不開放林農自行修改／刪除已回報之明細項。若回報錯誤，請"
    "聯絡計畫團隊（admin）協助修正。未來版本將提供林農自助編輯功能。")

add_heading(doc, "(四) Q：什麼時候要結案？", 3)
add_para(doc,
    "A：當您判斷本許可案件所有葉片採收已回報完畢、不再新增時即可結案。建議於"
    "效期屆滿前完成結案。結案後不能再回報、不能編輯，狀態固定供分署查核與合作"
    "社共同銷售彙整。")

add_heading(doc, "四、公文與許可單相關", 2)

add_heading(doc, "(一) Q：申請公文稿一定要列印嗎？", 3)
add_para(doc,
    "A：申請於本系統送出時，分署端會即時於審核分頁看到該案件。但**仍以紙本"
    "公文為法定文書**，請列印申請公文稿並親簽蓋章後，依分署規定方式（紙本郵寄、"
    "現場遞交、或掃描回傳 email）送達分署承辦人。")

add_heading(doc, "(二) Q：許可單需要隨身攜帶嗎？", 3)
add_para(doc,
    "A：建議於修枝作業期間隨身攜帶採取許可單列印版作為作業合法性證明。"
    "巡護人員、查緝人員可能要求出示許可單。")

add_heading(doc, "(三) Q：列印時上方有按鈕列怎麼辦？", 3)
add_para(doc,
    "A：列印頁面上方之「🖨️ 列印 / ✕ 關閉並返回」按鈕列為灰底設計、且套用 "
    "@media print 規則，實際列印或存 PDF 時會自動隱藏不出現於輸出文件。請放心"
    "直接按瀏覽器列印功能。")

add_heading(doc, "五、系統技術相關", 2)

add_heading(doc, "(一) Q：頁面卡住不動怎麼辦？", 3)
add_para(doc,
    "A：請先按 Ctrl + Shift + R（Windows）或 Cmd + Shift + R（Mac）強制重新"
    "整理。若仍無法，請開啟新分頁、清除瀏覽器資料後重試。若問題持續，請聯絡"
    "計畫團隊技術窗口並提供：(a) 瀏覽器類型與版本 (b) 操作步驟 (c) 螢幕截圖。")

add_heading(doc, "(二) Q：可以離線使用嗎？", 3)
add_para(doc,
    "A：本系統為 PWA 設計，可離線開啟主畫面與既有資料（Firestore offline "
    "persistence）。但送出申請、核准、回報採收量等寫入操作需要連網，離線時"
    "暫存、回連時自動同步。野外無訊號區域操作前，請先於有訊號處完整登入過一次。")

add_heading(doc, "(三) Q：哪些瀏覽器可以用？", 3)
add_para(doc,
    "A：建議使用 Chrome 或 Edge（最佳支援）。Firefox、Safari 亦可使用，但部分"
    "進階功能（PWA 安裝、相機 AI 辨識）可能受限。請避免使用 IE 11 或更舊版本。")

# ─────────────────────────────────────────────────────
# 附錄
# ─────────────────────────────────────────────────────
add_heading(doc, "附錄", 1)

add_heading(doc, "一、詞彙對照表", 2)
glossary = [
    ["系統內名稱", "中文說明", "權限類別"],
    ["admin / isSystemAdmin", "系統管理員（總管理者）", "頂層"],
    ["pi", "計畫主持人", "專案層"],
    ["surveyor", "調查員（本系統中為林農）", "專案層"],
    ["harvest_authority", "採取許可審核人員（林業保育署承辦）", "專案層"],
    ["coop", "林業合作社（彙整觀察者）", "專案層"],
    ["harvestPermit", "修枝申請（含葉片採收）", "資料表"],
    ["landParcelMeta", "申請單地籍結構化 metadata", "欄位"],
    ["boundaryGeoJsonStr", "專案邊界（FeatureCollection JSON 字串）", "欄位"],
    ["applicantSeq / applicantDocNo", "發文字號流水與字號（合作社編）", "欄位"],
    ["permitSeq / permitNo", "法定許可文號流水與字號（分署編）", "欄位"],
    ["totalLogged_kg", "葉片採收回報累計（denormalized）", "欄位"],
]
add_table_with_header(doc, glossary[0], glossary[1:], col_widths=[5.5, 6.0, 3.0])

add_heading(doc, "二、本系統使用相關法源", 2)
add_para(doc, "1. 森林法（中華民國 105 年 11 月 30 日修正版）")
add_para(doc, "2. 國有林林產物處分規則")
add_para(doc, "3. 大雪山林業合作社與林業保育署臺中分署土肉桂專區契約附圖（合作社契約）")
add_para(doc, "4. IPCC 2006 Guidelines for National Greenhouse Gas Inventories (LULUCF)")

add_heading(doc, "三、聯絡窗口", 2)
add_para(doc,
    "系統技術問題：計畫團隊（陳朝圳教授）。"
    "土肉桂專區契約與作業協調：合作社主席。"
    "修枝許可審核：林業保育署臺中分署承辦人員。")

add_heading(doc, "四、版本對應", 2)
add_para(doc,
    "本手冊對應 ForestMRV prod v2.11.69（2026 年 5 月 30 日 ship）。系統未來版本"
    "若有重大功能調整將另發補充說明文件。可於主畫面左上角版本徽章確認當前版本，"
    "若徽章版本高於本手冊，請參閱 service-worker.js 內嵌 changelog 或聯絡計畫"
    "團隊。")

# 存檔
doc.save(OUT)
print(f"OK: {OUT}")
print(f"Size: {OUT.stat().st_size / 1024:.1f} KB")
