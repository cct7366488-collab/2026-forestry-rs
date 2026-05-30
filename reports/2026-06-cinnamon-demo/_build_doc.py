"""
產生：01-角色矩陣與演示腳本.docx
為 6/6 土肉桂專區系統說明會準備（v2.11.69）。

依 CLAUDE.md 全域排版規則：
  - 5 級標號：壹、 / 一、 / (一) / 1. / (1)
  - 段落首行縮排 2 字（半形空格不可用、用首行縮排屬性）
  - 無項目符號（編號清單）
  - 無全形空格（U+3000）
  - CJK 字型：宋體（Times New Roman 西文 + DFKai-SB 中文 / PMingLiU）

跑法：python _build_doc.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).parent / "01-角色矩陣與演示腳本.docx"

# ===== 樣式設定 =====
CJK_FONT = "PMingLiU"   # 新細明體；可改 "DFKai-SB" 標楷體
ASCII_FONT = "Times New Roman"
INDENT_2CHAR_TWIPS = 480   # 約等於 2 個中文字（12pt × 2 × 20 twips）


def set_cjk_font(run, name=CJK_FONT, ascii_name=ASCII_FONT, size=12):
    """設定 run 的 CJK + ASCII 字型 + 字級。"""
    run.font.name = ascii_name
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), ascii_name)
    rFonts.set(qn('w:hAnsi'), ascii_name)


def add_para(doc, text, size=12, bold=False, indent=True, color=None, align=None, after_pt=4):
    """加段落。indent=True 套首行縮排 2 字。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after_pt)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.85)   # ~2 個 12pt 中文字
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    set_cjk_font(run, size=size)
    return p


def add_heading(doc, text, level, after_pt=6):
    """加標號（壹、一、(一)、1.、(1)）— 不用 Word 內建 Heading 樣式以保格式控制。"""
    sizes = {1: 16, 2: 14, 3: 13, 4: 12, 5: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(after_pt)
    if level >= 3:
        p.paragraph_format.first_line_indent = Cm(0.85)
    run = p.add_run(text)
    run.bold = True
    set_cjk_font(run, size=sizes.get(level, 12))
    return p


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    set_cjk_font(run, size=18)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(text)
    set_cjk_font(run, size=12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_table_with_header(doc, headers, rows, col_widths=None):
    """加表格，第一列為表頭（粗體底色）。"""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    # 表頭
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        set_cjk_font(run, size=11)
        # 表頭底色淺藍
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'DCE6F1')
        tcPr.append(shd)
    # 內容列
    for r_i, row in enumerate(rows, 1):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i].cells[c_i]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_cjk_font(run, size=10)
    if col_widths:
        for col, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[col].width = Cm(w)


# ===== 開始建文件 =====
doc = Document()

# 頁面邊距
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

add_title(doc, "ForestMRV 土肉桂專區系統")
add_title(doc, "角色矩陣與演示腳本")
add_subtitle(doc, "6/6 系統說明會準備文件 │ v2.11.69 │ 2026-05-30")

# ─────────────────────────────────────────────────────
# 壹、文件目的與系統定位
# ─────────────────────────────────────────────────────
add_heading(doc, "壹、文件目的與系統定位", 1)
add_para(doc,
    "本文件為 2026 年 6 月 6 日土肉桂專區葉片收穫監測會計系統說明會準備。"
    "對象包含林業保育署臺中分署承辦人員、土肉桂專區合作社代表、契約林農、"
    "以及計畫團隊內部成員。文件涵蓋五個系統角色之權責矩陣、五階段操作流程、"
    "30 分鐘說明會演示腳本，以及目前已知議題與後續工作。")

add_heading(doc, "一、系統定位與計畫脈絡", 2)
add_para(doc,
    "ForestMRV 為依循 IPCC LULUCF 與 MRV（測量、報告、查證）原則建構的森林"
    "監測平臺。針對土肉桂專區，本系統承接以下行政流程：林農申請修枝採取"
    "副產物（葉片）→ 林業保育署臺中分署核准並核發法定許可文號 → 林農回報"
    "實際葉片採收量 → 結案後資料移交合作社進行共同銷售彙整。")

add_heading(doc, "二、現階段與未來移交", 2)
add_para(doc,
    "現階段（計畫執行期）系統管理員與計畫主持人皆由計畫團隊負責人擔任。"
    "計畫結束移交後，系統管理員仍由計畫團隊保留以維護資料完整性；計畫主"
    "持人將移交予合作社主席，由合作社接手後續日常營運。林農、審核人員之"
    "角色設計於兩階段皆保持一致。")

add_heading(doc, "三、本版（v2.11.69）相對 5/20 demo 的關鍵變動概覽", 2)
add_para(doc,
    "自 5/20 stakeholder 說明會（v2.11.40）後，系統依各方回饋陸續釋出 11 個版本，"
    "於 5/30 dry-run 完成最終 polish 並 ship v2.11.68 / v2.11.69。對 6/6 demo 觀眾"
    "最關鍵的差異有四：")
add_heading(doc, "(一) 修枝申請主秀大幅升級", 3)
add_para(doc,
    "三層連動下拉（v2.11.55）→ 申請人姓名自動帶承租人（v2.11.63）→ 株數依面積"
    "×1800 株/ha 自動計算（v2.11.59）→ 表單即時嵌入 Leaflet 迷你地圖（v2.11.66）→"
    "申請函嵌入自製 SVG 作業位置示意圖（v2.11.67）。林農「先選作業位置 → 各欄"
    "位自動帶入」的順序由 v2.11.68 表單欄位重排定案。")
add_heading(doc, "(二) 公文整合終章", 3)
add_para(doc,
    "申請函說明改為公文體 4 點（v2.11.60）；發文字號改 Firestore runTransaction "
    "原子產生「大雪山林業合作社（修）字第 NNN 號」確保不重號（v2.11.61）；"
    "法定許可文號亦同機制產出「林保中-土肉桂修枝-115-NNN」。電子＋紙本雙軌即"
    "時同產，紙本內容由線上記錄產生、不漂移。")
add_heading(doc, "(三) 純行政專案介面收斂", 3)
add_para(doc,
    "土肉桂專區屬純行政（非樣區調查）專案，v2.11.64 隱藏不需要的樣區/設計/儀"
    "表板分頁，v2.11.68 補強地圖頁的立木相關控制元件（著色依據／樣區邊界圖層）"
    "亦於純行政模式藏起，介面更貼合實際業務需求。")
add_heading(doc, "(四) 業務拍板：許可不預估採收量", 3)
add_para(doc,
    "v2.11.69 砍掉「核准葉片採收量 (kg)」概念。理由：土肉桂每塊地年齡不同、每株"
    "產出葉量差異大，事先預估上限反而誤導承辦與林農。改為「許可只許可修枝作業，"
    "所採取葉片重量以林農實際回報為準、不設預估上限」並於採取許可單明文宣告。"
    "核准 modal 由原本三欄位（核准量 + 效期起 + 效期迄）簡化為單欄位（效期），承辦"
    "人作業負擔降低；合作社彙整僅顯示「林農／案件數／已回報葉片採收量(kg)」三欄"
    "純累計、不對照、不算達成率。")

# ─────────────────────────────────────────────────────
# 貳、角色與權責矩陣
# ─────────────────────────────────────────────────────
add_heading(doc, "貳、角色與權責矩陣", 1)

add_heading(doc, "一、角色定義", 2)
role_def = [
    ["角色（系統名）", "現階段擔任", "未來移交對象", "核心職責"],
    ["系統管理員（admin）", "計畫團隊", "不變", "開案、設邊界、人員指派、QAQC 監督"],
    ["計畫主持人（pi）", "計畫團隊指派", "合作社主席", "方法學設定、樣區規格、邀請成員"],
    ["林農（surveyor）", "契約承租林農", "不變", "修枝申請、實際葉片採收量回報、結案"],
    ["審核人員（harvest_authority）", "林業保育署臺中分署承辦", "不變", "核准／駁回／要求補件、核發法定許可文號"],
    ["合作社（coop）", "合作社主席", "升級為計畫主持人", "唯讀彙整檢視、共同銷售匯出 Excel"],
]
add_table_with_header(doc, role_def[0], role_def[1:], col_widths=[3.0, 3.0, 3.0, 6.5])

add_heading(doc, "二、權限矩陣對照表", 2)
add_para(doc,
    "下表列示五個角色於系統七項主要功能之權限對照。「✅」表示具完整存取與"
    "操作權；「⚠️」表示部分受限或唯讀；「❌」表示無存取權。權限以"
    "Firestore Security Rules 強制執行，介面層 UI gating 為輔助。")

perm_matrix = [
    ["功能", "admin", "pi", "surveyor", "harvest_authority", "coop"],
    ["專案開案", "✅", "❌", "❌", "❌", "❌"],
    ["上傳／更新專案邊界", "✅", "✅", "❌", "❌", "❌"],
    ["方法學設定", "✅", "✅", "❌", "❌", "❌"],
    ["新增樣區與立木", "✅", "✅", "✅（限派任）", "❌", "❌"],
    ["提出修枝申請", "❌", "✅", "✅", "❌", "❌"],
    ["審核修枝申請", "✅", "❌", "❌", "✅", "❌"],
    ["填報葉片採收量", "❌", "✅", "✅（限本人）", "❌", "❌"],
    ["合作社彙整檢視", "✅", "✅", "❌", "❌", "✅"],
    ["匯出 Excel", "✅", "✅", "❌", "❌", "✅"],
    ["刪除申請案（任何狀態）", "✅", "✅", "⚠️（僅自己草稿）", "❌", "❌"],
]
add_table_with_header(doc, perm_matrix[0], perm_matrix[1:], col_widths=[4.5, 1.8, 1.8, 1.8, 2.5, 1.8])

add_heading(doc, "三、設計原則說明", 2)
add_heading(doc, "(一) 邊界資料受法律約束", 3)
add_para(doc,
    "土肉桂專區作業單元邊界源自合作社與林業保育署簽訂之契約附圖，具法律"
    "約束力。系統設計允許 admin 與 pi 修改邊界，但於 6 月 6 日說明會後將"
    "視 PI 移交合作社主席之時程，再評估是否將邊界寫入權限收緊至 admin only。")

add_heading(doc, "(二) 林農權限單一化", 3)
add_para(doc,
    "林農（surveyor）角色於本系統內僅負責申請與回報，不得修改邊界、不得"
    "審核他人申請、不得跨案讀寫。同一林農於不同專案理論上可被指派不同角"
    "色，但於土肉桂專區內角色固定為 surveyor。")

add_heading(doc, "(三) 合作社唯讀並具未來升級路徑", 3)
add_para(doc,
    "合作社現階段為唯讀彙整角色，主要用於共同銷售前彙整各林農已結案之葉"
    "片採收量並匯出 Excel。下一輪開發（v2.11.70+）將開放合作社對申請單之留"
    "言註記功能（無否決權），協助合作社與林農溝通採收時程或品質要求。未來"
    "移交後合作社主席將升級為 pi 角色，接手日常營運。")

add_heading(doc, "(四) PI／admin 可跨狀態刪除案件", 3)
add_para(doc,
    "v2.11.62 新增 PI／admin 對申請案於任何狀態（草稿／送出／審核中／已核准／"
    "修枝作業中／已結案／已駁回）皆可刪除之後門，cascade 級聯刪除子集合 logs。"
    "確認對話框會列出申請人／林地／狀態／發文字號／法定許可文號（警示將同時失"
    "效）／累計回報量（警示子集合一併刪除）。林農 owner 仍僅限於草稿狀態刪自己"
    "的案。設計目的：管理員於 demo 後可一鍵清掉測試資料；正式啟用後亦保留管理"
    "後門以處理誤送、重送、需作廢之案件。")

# ─────────────────────────────────────────────────────
# 參、目前系統已實作功能（v2.11.69）
# ─────────────────────────────────────────────────────
add_heading(doc, "參、目前系統已實作功能（v2.11.69）", 1)

add_heading(doc, "一、5 月 20 日說明會後重要更新總覽", 2)
add_para(doc,
    "自 v2.11.40（5/20 demo 版）起共釋出 29 個版本至 v2.11.69（6/6 demo 版）。"
    "下表列出對 6/6 demo 觀眾最有感的關鍵版本。詳細歷程請參閱 git log 與 SW "
    "service-worker.js 內嵌 changelog 註解。")
update_log = [
    ["版本", "主要更新", "對應角色"],
    ["v2.11.40 (5/20)", "B1 文案調整：申請主體＝「修枝」、事後實際量＝「葉片採收」", "全角色"],
    ["v2.11.41", "申請表單精簡：移除預估產出量與用途欄位", "surveyor"],
    ["v2.11.52", "樣區上傳防呆：偵測誤將專案邊界當樣區上傳", "admin"],
    ["v2.11.53", "編輯專案雙入口：頂部按鈕＋設定分頁區塊", "admin / pi"],
    ["v2.11.54", "開案表單支援邊界上傳＋預設邊界一鍵載入", "admin"],
    ["v2.11.55", "修枝申請三層連動下拉：專區→承租人→作業單元，自動帶出地籍", "surveyor"],
    ["v2.11.58", "修申請面積欄位 step 精度：允許 4 位小數（1 m² 精度）", "surveyor"],
    ["v2.11.59", "株數＝面積×1800 株/ha 自動連動；修枝起迄日改必填", "surveyor"],
    ["v2.11.60", "申請函說明改公文體 4 點（撫育／原則／詳如表列／系統聲明）", "surveyor + 承辦"],
    ["v2.11.61", "發文字號 runTransaction 原子產生「大雪山林業合作社（修）字第 NNN 號」", "surveyor"],
    ["v2.11.62", "PI／admin 可刪除申請案任何狀態（cascade logs）", "admin / pi"],
    ["v2.11.63", "申請人姓名 picker 自動帶承租人；聯絡方式必填", "surveyor"],
    ["v2.11.64", "純行政專案隱藏 樣區／設計／儀表板 三分頁", "admin"],
    ["v2.11.65", "主地圖 onEachFeature 加 label + popup（115 個作業單元代碼）", "admin / pi / coop"],
    ["v2.11.66", "申請表單嵌 240px Leaflet 迷你地圖（picker 即時預覽）", "surveyor"],
    ["v2.11.67", "申請函嵌入自製 SVG 作業位置示意圖（紅目標＋灰 context＋N 箭頭）", "surveyor + 承辦"],
    ["v2.11.68 (5/30)", "純行政專案地圖頁清掉立木 UI 殘留；修枝申請欄位重排", "admin / surveyor"],
    ["v2.11.69 (5/30)", "拿掉「核准採收量」概念：許可不預估上限、採收量以實際回報為準", "全角色"],
]
add_table_with_header(doc, update_log[0], update_log[1:], col_widths=[2.5, 9.0, 3.5])

add_heading(doc, "二、申請帶出地籍的演示亮點（v2.11.55 → v2.11.69 累積）", 2)
add_para(doc,
    "本次說明會主秀為土肉桂修枝申請流程。林農於修枝申請表單依下列順序自然完成"
    "作業位置選擇與資料填報，全程一次到位、不需離開表單檢索其他資料：")

add_heading(doc, "(一) 第一層 — 專區（v2.11.55）", 3)
add_para(doc,
    "下拉選單列出本專案邊界涵蓋之專區清單。土肉桂專區包含「橫流溪」與「烏石坑」"
    "兩個地理區域。")

add_heading(doc, "(二) 第二層 — 承租人（v2.11.55 / v2.11.63）", 3)
add_para(doc,
    "依第一層所選之專區自動篩選該區所有契約承租人。系統依姓名字典序排列。"
    "v2.11.63 加強：承租人一旦選定即同步自動覆蓋上方「申請人姓名」欄位（承租人"
    "＝申請人為原則）；user 仍可手動覆蓋此自動帶入值。")

add_heading(doc, "(三) 第三層 — 作業單元（v2.11.55）", 3)
add_para(doc,
    "依第二層所選之承租人列出其名下之所有作業單元。每筆顯示作業單元編號"
    "（如「14_1_1」或「117-107-01」）、契約面積（ha）、契約樹種摘要。")

add_heading(doc, "(四) 自動帶入欄位（v2.11.55 + v2.11.59）", 3)
add_para(doc,
    "完成三層選擇後，系統自動帶入「林班 / 假地號」與「申請修枝面積」兩個欄位"
    "（仍可手動編輯覆蓋）。v2.11.59 新增「土肉桂修枝株數」依面積 × 1800 株/ha"
    "（合作社契約假設密度）自動計算填入，林農修改面積時即時連動更新株數；亦可"
    "手動覆蓋此預設值。同時於下方顯示「✓ 已帶入下方林班 / 地號欄位」之確認摘要，"
    "含 7 欄結構化資料：專區、承租人、作業單元、契約樹種、契約面積、工作站 / 事業"
    "區、契約書號。")

add_heading(doc, "(五) 即時地圖預覽（v2.11.66）", 3)
add_para(doc,
    "申請表單下方嵌入 240px 高 Leaflet 迷你地圖。初始顯示專案全部 115 個作業單元"
    "（灰色淡填、fit 全 boundary）；picker 選定承租人後，紅色高亮該承租人名下所有"
    "圖塊並 fit 對應 bounds；picker 完成至作業單元後，紅色高亮該單一作業單元並"
    "fit（maxZoom 18）。下方 status 行即時更新「✓ 已定位：橫流溪 / 陳世允 / 14_1_1"
    "（1 個圖塊）」之類訊息。林農無須切換到地圖頁即可確認所選作業單元位置正確。")

add_heading(doc, "(六) 結構化資料保留（v2.11.55）", 3)
add_para(doc,
    "提交申請時，三層選擇結果以結構化 metadata 儲存至 Firestore（欄位名為 "
    "landParcelMeta），供後續合作社彙整、許可單溯源、報表分析使用。申請卡片亦"
    "顯示「✓ 系統識別」摘要行，協助林農目視確認未選錯。")

add_heading(doc, "三、公文整合終章（v2.11.60 / 61 / 67）", 2)

add_heading(doc, "(一) 申請函說明改公文體 4 點（v2.11.60）", 3)
add_para(doc,
    "原 6 點逐項條列（與下表重複、流水帳）→ 4 點正式公文風格：(1) 撫育管理＋通風"
    "透光採枝葉理由 (2) 不損生長／不影響保育／避免過度修剪原則 (3) 範圍／數量／"
    "期間詳如本函表列資料 (4) ForestMRV 線上系統登錄聲明（草稿／送出兩種文案）。"
    "表列資料維持原樣。")

add_heading(doc, "(二) 發文字號 runTransaction 自動編號（v2.11.61）", 3)
add_para(doc,
    "申請函發文字號採 Firestore runTransaction 原子遞增 counters/applicantSeq 子文件，"
    "格式「大雪山林業合作社（修）字第 NNN 號」（三位零填補，自 001 開始）。林農送"
    "出申請瞬間配發、寫入 harvestPermit 文件供下次列印取用；已配發過（revision／"
    "再送／編輯）不重新編號，避免改卷時改文號破壞紙本連續性。")

add_heading(doc, "(三) SVG 作業位置示意圖嵌入申請函（v2.11.67）", 3)
add_para(doc,
    "申請函附件區自動嵌入自製 inline SVG 作業位置示意圖（無 leaflet-image／"
    "html2canvas 外部依賴、無 CORS、黑白印表友善、確定性高）。讀 project."
    "boundaryGeoJsonStr 找 meta.unitId 對應 polygon →收集 1.5× target bbox 範"
    "圍鄰近作業單元為 context →緯度補償 lon×cos(lat) 保留正確 aspect → 等比例"
    "縮放置中 420×300。樣式：目標紅（#fecaca 填／#991b1b 邊）、context 灰、N 北"
    "方箭頭右上、目標 unitId 粗體紅標籤、鄰近單元小灰標籤。CSS page-break-inside:"
    "avoid 確保附圖不被分頁切開。承辦人查閱申請時，圖文一致、位置一目了然。")

add_heading(doc, "(四) 法定許可文號 runTransaction（同 v2.11.34 P1 vertical slice）", 3)
add_para(doc,
    "分署承辦人核准申請時，系統以同樣的 Firestore runTransaction 原子遞增"
    "counters/harvestPermit 子文件，產出法定許可文號「林保中-土肉桂修枝-{民國年}-"
    "{NNN}」（如「林保中-土肉桂修枝-115-001」）。發文字號（合作社編）與法定許可"
    "文號（分署編）為兩個獨立編號系統，分別追蹤申請端與核准端流水序。")

# ─────────────────────────────────────────────────────
# 肆、6/6 演示腳本（30 分鐘）
# ─────────────────────────────────────────────────────
add_heading(doc, "肆、6 月 6 日演示腳本（30 分鐘）", 1)

add_heading(doc, "一、演示前置作業", 2)
add_para(doc,
    "建議演示時準備五個瀏覽器帳號（或一台筆電多分頁，登入不同帳號），分別代表"
    "系統管理員、計畫主持人、林農、審核人員、合作社五個角色。若無法準備五帳號，"
    "可採用簡化版：以系統管理員一個帳號穿越各分頁示範（admin 權限涵蓋全部功能），"
    "於切換場景時口頭說明「現在切換為某某角色」。")
add_para(doc,
    "現場操作時，建議瀏覽器使用無痕視窗模式並停用所有 Chrome 擴充功能，避免擴充"
    "功能於 DevTools Console 製造紅字（5/30 dry-run 確認皆為擴充功能 listener "
    "noise，與本系統無關）。若會中需開啟 DevTools 展示資料流向，務必先關閉擴充。")

add_heading(doc, "二、演示流程分段", 2)

add_heading(doc, "(一) 第 1 段（0-5 分鐘）── 系統概觀與角色介紹", 3)
add_para(doc, "1. 開啟 https://forestry-rs-monitor.web.app 主畫面、確認版本徽章為 v2.11.69。")
add_para(doc, "2. 簡介五個角色及其分工，引用本文件「貳、角色與權責矩陣」之表格。")
add_para(doc, "3. 進入「土肉桂專區葉片收穫監測、會計系統」專案首頁，介紹分頁佈局；")
add_para(doc, "   說明純行政專案僅顯示「地圖／修枝申請／修枝審核／葉片採收回報及結案／")
add_para(doc, "   葉片採收彙整／匯出／設定」七個分頁（樣區／設計／儀表板於 v2.11.64 已隱）。")

add_heading(doc, "(二) 第 2 段（5-10 分鐘）── 系統管理員開案與邊界準備", 3)
add_para(doc, "1. 切換至系統管理員視角，回到「我的專案」首頁。")
add_para(doc, "2. 點「＋ 新專案」展示開案表單；填入示範資料但不送出。")
add_para(doc, "3. 重點演示「📐 專案邊界（GeoJSON，選填）」區塊內的「📥 預設邊界」下拉。")
add_para(doc, "4. 選擇「土肉桂專區（橫流溪 + 烏石坑，合作社契約）」後按「載入」，展示綠字確認摘要。")
add_para(doc, "5. 說明此預設集為合併檔（橫流溪 45 + 烏石坑 70 = 115 個作業單元），含完整契約 metadata。")
add_para(doc, "6. 切到「地圖」分頁，展示 115 個作業單元邊界疊圖；v2.11.65 加入的「作業單元代碼」")
add_para(doc, "   tooltip label 即時顯示，點任一面跳出 7 欄結構化資訊 popup（作業單元／專區／")
add_para(doc, "   承租人／契約樹種／契約面積／工作站／契約書號）。")

add_heading(doc, "(三) 第 3 段（10-15 分鐘）── 林農修枝申請（演示主秀）", 3)
add_para(doc, "1. 切換至林農視角。")
add_para(doc, "2. 點「🌿 修枝申請」分頁，點「＋ 修枝申請」開啟表單。")
add_para(doc, "3. 重點演示三層連動下拉：")
add_para(doc, "   • 第一層 — 選「橫流溪」")
add_para(doc, "   • 第二層 — 系統列出該專區承租人，選任一人；申請人姓名欄即自動覆蓋為承租人姓名")
add_para(doc, "   • 第三層 — 系統列出該承租人作業單元，選任一筆")
add_para(doc, "4. 完成三層選擇後展示自動帶入之「林班 / 地號」、「申請修枝面積」與「土肉桂修枝株數」")
add_para(doc, "   三欄；強調株數係依面積×1800 株/ha 自動計算（v2.11.59）。")
add_para(doc, "5. 展示下方表單嵌入 Leaflet 迷你地圖：紅色高亮所選作業單元、自動 zoom（v2.11.66）。")
add_para(doc, "6. 展示 7 欄結構化摘要「✓ 已帶入下方林班 / 地號欄位」。")
add_para(doc, "7. 補填聯絡方式（必填）、修枝方式、修枝起訖日（必填）、備註後送出申請。")
add_para(doc, "8. 系統以 runTransaction 配發「大雪山林業合作社（修）字第 NNN 號」發文字號（v2.11.61）；")
add_para(doc, "   toast 顯示文號、卡片狀態變「已送出・待審」、出現「📜 發文字號」資訊行。")
add_para(doc, "9. 點「📄 申請公文稿」展示自動產出之公文「函」格式：")
add_para(doc, "   • 公文體 4 點說明（v2.11.60）")
add_para(doc, "   • 表列 8 欄申請明細")
add_para(doc, "   • 附件嵌入 SVG 作業位置示意圖（紅目標＋灰 context＋N 箭頭、v2.11.67）")
add_para(doc, "   • 申請人簽章區 + 分署收件欄")

add_heading(doc, "(四) 第 4 段（15-20 分鐘）── 審核人員核准", 3)
add_para(doc, "1. 切換至審核人員（林業保育署臺中分署承辦人）視角。")
add_para(doc, "2. 點「📋 修枝審核」分頁，展示待審清單及紅點提示；副標「審核者僅能核准／要求補件／")
add_para(doc, "   駁回，不可變更申請內容（權限由 Firestore 規則鎖定）」說明分權設計。")
add_para(doc, "3. 開啟剛才送出之申請，演示「核准」流程：")
add_para(doc, "   • 系統顯示申請人摘要（read-only）")
add_para(doc, "   • 灰底提示框宣告「本案核准後即許可進行修枝作業；所採取葉片重量由林農於 ForestMRV")
add_para(doc, "     線上系統實際回報，不設預估上限」（v2.11.69 拍板）")
add_para(doc, "   • 承辦人僅需填入效期起日／迄日，與選填審核附註")
add_para(doc, "4. 確認送出後系統以 runTransaction 原子產生法定許可文號「林保中-土肉桂修枝-115-NNN」；")
add_para(doc, "   toast 顯示文號、待審清單數量 -1、案件移入已處理區。")
add_para(doc, "5. 切換回林農視角，展示「採取許可單」按鈕及其列印出之正式許可格式：")
add_para(doc, "   • 公文 box 含許可文號／申請人／林地／修枝方式／效期 5 行")
add_para(doc, "   • 底部一句宣告「本許可只許可於上述林地進行土肉桂修枝作業；所採取葉片重量以申請人")
add_para(doc, "     於 ForestMRV 線上系統實際回報為準、不設預估上限」（v2.11.69 新增）")
add_para(doc, "   • 葉片採收量登錄紀錄表（核准時為空）")
add_para(doc, "   • 申請人簽章 + 分署核章欄")

add_heading(doc, "(五) 第 5 段（20-25 分鐘）── 林農回報與結案", 3)
add_para(doc, "1. 切到「🌾 葉片採收回報及結案」分頁。")
add_para(doc, "2. 點開剛核准之案件，演示「＋ 填報葉片採收量」（可分批多次回報）。")
add_para(doc, "3. 填入採收日期、實際鮮葉重(kg)，選填乾燥後重、含水率、批次、備註。送出後 toast")
add_para(doc, "   顯示「✅ 已回報，累計 X kg」（v2.11.69 起不再做核准量對照、無達成率／超量警示）。")
add_para(doc, "4. 系統首筆回報時自動推進狀態 approved → harvesting；卡片即時顯示「已回報葉片採收")
add_para(doc, "   累計：X kg」與明細表。")
add_para(doc, "5. 全數回報完成後按「✅ 回報完畢並結案」結束本筆案件；")
add_para(doc, "   confirm 對話框列出「累計回報葉片採收量：X kg」（不再對照核准量、v2.11.69 簡化）。")
add_para(doc, "6. 結案閘門：rules 與 client 雙擋零回報結案（必須 totalLogged_kg > 0 才能 → completed）。")

add_heading(doc, "(六) 第 6 段（25-30 分鐘）── 合作社彙整與 Q&A", 3)
add_para(doc, "1. 切換至合作社視角，點「📊 葉片採收彙整」分頁。")
add_para(doc, "2. 展示三個區塊：")
add_para(doc, "   • 📋 專區總覽：申請案件數、狀態分布、已回報葉片採收量總計")
add_para(doc, "   • 🤝 依林農葉片採收彙整（共同銷售）：三欄表「林農／案件數／已回報葉片採收量(kg)」")
add_para(doc, "     依採收量降序排列；無達成率欄（v2.11.41 拍板）")
add_para(doc, "   • 📋 各申請案（唯讀）：每筆顯示申請人／狀態／文號／林地／已回報量／採取許可單按鈕")
add_para(doc, "3. 演示「📥 匯出 Excel」供共同銷售前資料整理。")
add_para(doc, "4. 預留 Q&A 時間。")

# ─────────────────────────────────────────────────────
# 伍、已知議題與後續工作
# ─────────────────────────────────────────────────────
add_heading(doc, "伍、已知議題與後續工作", 1)

add_heading(doc, "一、5/30 dry-run 後已修補完成（v2.11.68 / v2.11.69 ship）", 2)
done_log = [
    ["代號", "項目", "完成版本"],
    ["F1", "純行政專案地圖頁立木 UI 殘留（著色依據 radio + 樣區邊界 checkbox）", "v2.11.68"],
    ["F5", "修枝申請表單欄位順序（picker / mini-map 移到最上）", "v2.11.68"],
    ["F6 → 業務拍板", "拿掉「核准採收量」概念（許可不預估上限、實際回報為準）", "v2.11.69"],
]
add_table_with_header(doc, done_log[0], done_log[1:], col_widths=[2.0, 9.5, 3.5])

add_heading(doc, "二、6/6 後規劃之開發工作", 2)
todo_list = [
    ["優先序", "代號", "工作項目", "規模", "建議完成時程"],
    ["P0", "—", "合作社（coop）對申請單留言註記功能（無否決權）", "1-2 小時", "v2.11.70（6 月中）"],
    ["P1", "—", "邊界圖檔自動套疊地籍圖以驗證契約面積", "0.5 天", "6 月底"],
    ["P2", "—", "未來移交時 admin 收緊邊界寫入權至 admin only", "需與合作社協調", "PI 移交前"],
    ["P3", "—", "預設邊界檔從 Hosting 靜態檔遷至 Firestore boundaryPresets 集合", "1 天", "視 PII 風險評估"],
    ["P4", "F3", "Tailwind CDN → CLI build / PostCSS 整合", "1 小時", "下季度"],
    ["P5", "F4", "console.log 雜訊改為 verbose-only flag 控制", "15 分鐘", "下季度"],
    ["P6", "—", "其他未對照計畫類型補套餐（TRAIN、公園碳匯等）", "1 天", "下季度"],
]
add_table_with_header(doc, todo_list[0], todo_list[1:], col_widths=[1.5, 1.5, 7.5, 2.5, 2.5])

add_heading(doc, "三、已知技術限制與風險", 2)

add_heading(doc, "(一) 預設邊界檔含承租人姓名 PII", 3)
add_para(doc,
    "目前合併檔放置於 Hosting 公開靜態目錄，URL 已知者皆可下載。判定為合"
    "作社契約半公開資料（林業保育署本身有開放類似 dataset）、demo 演示風"
    "險可接受。若 6/6 後利害關係人對 PII 保護有疑慮，可遷至 Firestore 並"
    "設 rules 鎖 admin 才能讀。")

add_heading(doc, "(二) 烏石坑檔案部分欄位名 UTF-8 截斷", 3)
add_para(doc,
    "原始檔案（烏石坑土肉桂作業單元.geojson）「契約樹種」「伐木作業」「造"
    "林作業」「撫育作業」「契約面積」等欄位名於 export 過程中末字組節被"
    "截斷。合併腳本已自動修補，但建議下次匯出時於來源端使用 UTF-8 嚴格"
    "模式檢查避免汙染。")

add_heading(doc, "(三) 既有舊版邊界需手動升級一次", 3)
add_para(doc,
    "v2.11.54 以前上傳之邊界僅保留 geometry、未保留 properties。升級至"
    "v2.11.55 後，須由 admin 進入「✏️ 編輯專案 → 📐 專案邊界 → 📥 預"
    "設邊界 → 載入 → 儲存」一次性升級為含 metadata 之 FeatureCollection。"
    "此手動步驟僅需做一次，後續上傳即自動為新格式。")

add_heading(doc, "(四) Demo 機 Chrome 擴充功能 Console 噪音", 3)
add_para(doc,
    "5/30 dry-run 確認 Console 紅字「A listener indicated an asynchronous "
    "response by returning true, but the message channel closed」為 Chrome "
    "擴充功能造成、非系統 bug。demo 機建議使用無痕視窗 + 停用全部擴充功能"
    "以避免 DevTools 展示時的視覺干擾。")

add_heading(doc, "(五) approvedAmount_kg 既有資料保留", 3)
add_para(doc,
    "v2.11.69 砍掉「核准採收量」概念後，既有 Firestore 內已存兩筆測試"
    "資料（吳宏斌 null kg、陳璽元 0.9 kg）approvedAmount_kg 欄位仍保留。UI 已"
    "全面隱藏不顯示，業務上不影響；schema 保留欄位以利向後相容，若未來改變"
    "心意可恢復顯示。若 6/6 前希望徹底清理，可於 admin 端執行 Firestore migration"
    "腳本將欄位 unset。")

# ─────────────────────────────────────────────────────
# 附錄
# ─────────────────────────────────────────────────────
add_heading(doc, "附錄", 1)

add_heading(doc, "一、詞彙對照表", 2)
glossary = [
    ["系統內名稱", "中文說明", "權限類別"],
    ["admin / isSystemAdmin", "系統管理員（總管理者）", "頂層"],
    ["pi", "計畫主持人", "專案層"],
    ["surveyor", "調查員（本系統中為林農）", "專案層"],
    ["harvest_authority", "採取許可審核人員（林業保育署承辦）", "專案層"],
    ["coop", "林業合作社（彙整觀察者）", "專案層"],
    ["reviewer", "QAQC 查證員", "專案層"],
    ["harvestPermit", "修枝申請（含葉片採收）", "資料表"],
    ["landParcelMeta", "申請單地籍結構化 metadata", "欄位"],
    ["boundaryGeoJsonStr", "專案邊界（FeatureCollection JSON 字串）", "欄位"],
    ["applicantSeq / applicantDocNo", "發文字號流水與字號（合作社編）", "欄位"],
    ["permitSeq / permitNo", "法定許可文號流水與字號（分署編）", "欄位"],
    ["totalLogged_kg", "葉片採收回報累計（denormalized）", "欄位"],
    ["approvedAmount_kg", "核准採收量（v2.11.69 後 UI 全藏、寫 null）", "欄位（保留）"],
]
add_table_with_header(doc, glossary[0], glossary[1:], col_widths=[5.5, 6.0, 3.0])

add_heading(doc, "二、相關文件", 2)
add_para(doc,
    "本系統其他相關文件存放於 reports/2026-05-mid-demo/，包括：")
add_para(doc, "1. 01-簡報-ForestMRV 系統介紹.pptx：系統整體概念簡報")
add_para(doc, "2. 02-操作手冊-ForestMRV.docx：使用者操作詳細手冊")
add_para(doc, "3. 03-現場練習指南-下午 hands-on.docx：實機操作練習")
add_para(doc, "4. 04-土肉桂採收許可-520 說明會 demo 腳本.md：5/20 demo 原始腳本")
add_para(doc, "5. 05-土肉桂採收許可-520 說明會簡報.pptx：5/20 簡報")

add_heading(doc, "三、版本歷程與變動原則", 2)
add_para(doc,
    "完整 changelog 嵌於 service-worker.js 註解內（依版本倒序），可於該檔搜尋")
add_para(doc, "「v2.11.XX」快速定位該版主要改動。本說明會聚焦於 v2.11.40（5/20 demo")
add_para(doc, "版）至 v2.11.69（6/6 demo 版）之間的 29 個版本增量。版本變動原則：")
add_para(doc, "1. ?v= cache busting 全 15 檔 lockstep，避免 ESM 雙實例載入。")
add_para(doc, "2. SW CACHE 名稱依版本 bump，activate 時自動清除舊 cache。")
add_para(doc, "3. Schema 變動以向後相容為原則（既有資料不破壞，舊版讀新欄位視為缺值）。")
add_para(doc, "4. Rules 變動同步保留向後相容（舊欄位於白名單，新欄位加 validate）。")

add_heading(doc, "四、聯絡窗口", 2)
add_para(doc,
    "系統技術問題：計畫團隊（陳朝圳教授）。"
    "土肉桂專區契約與作業協調：合作社主席。"
    "修枝許可審核：林業保育署臺中分署承辦人員。")

# 存檔
doc.save(OUT)
print(f"OK: {OUT}")
print(f"Size: {OUT.stat().st_size / 1024:.1f} KB")
