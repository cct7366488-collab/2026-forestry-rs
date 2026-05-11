"""Retry H01 embed only (Word was holding file lock first time)."""
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Cm

BASE = Path(r"H:\我的雲端硬碟\2026 forestry_RS\reports\2026-05-mid-demo")
IMG = BASE / "_build" / "images"
H01 = IMG / "H01-章節導引示意.png"
HOL_PATH = BASE / "03-現場練習指南-下午hands-on.docx"
STAMP = datetime.now().strftime("%Y%m%d-%H%M%S")

bak = HOL_PATH.with_suffix(HOL_PATH.suffix + f".bak-{STAMP}")
shutil.copy2(HOL_PATH, bak)
print(f"backup -> {bak.name}")

doc = Document(str(HOL_PATH))
target_para = None
for p in doc.paragraphs:
    if "H01" in p.text and "章節導引" in p.text:
        target_para = p
        break
if target_para is None:
    raise SystemExit("HOL: H01 placeholder paragraph not found")

for r in list(target_para.runs):
    r._element.getparent().remove(r._element)
run = target_para.add_run()
run.add_picture(str(H01), width=Cm(14.0))
doc.save(str(HOL_PATH))
print("H01 picture inserted -> saved")
